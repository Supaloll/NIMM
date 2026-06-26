"""Génération de documents ACCESSIBLES à partir d'un contenu structuré.

Utilisé par l'outil CoaNIMM `nimm_make_document`. Produit des fichiers dont la
structure est lisible par un lecteur d'écran : un titre principal, des sous-titres
(headings), des paragraphes, et des images TOUJOURS accompagnées de leur description
(texte alternatif). Langue du document déclarée (utile au lecteur d'écran).

Entrée : `title` (str) + `sections` (liste de dicts). Chaque section peut porter :
  - 'titre' (str)  : un sous-titre (heading de niveau 1) ;
  - 'texte' (str)  : du corps (les paragraphes sont séparés par une ligne vide) ;
  - 'image' (str)  : chemin d'un fichier image à insérer ;
  - 'alt'   (str)  : description de l'image (obligatoire pour l'accessibilité).

Formats : 'html', 'docx', 'pdf', 'epub', 'txt'. Retour : (bytes, extension).
"""

import io
import os
import html as _html
import base64
import zipfile


def _norm_sections(sections):
    out = []
    for s in (sections or []):
        if isinstance(s, str):
            out.append({'texte': s})
        elif isinstance(s, dict):
            out.append(s)
    return out


def _paras(texte):
    """Découpe un corps en paragraphes (séparés par une ligne vide)."""
    if not texte:
        return []
    blocs = [p.strip() for p in str(texte).replace('\r\n', '\n').split('\n\n')]
    return [p for p in blocs if p]


def _img_data_uri(path):
    """Lit une image et retourne (data_uri, mime) ou (None, None)."""
    try:
        with open(path, 'rb') as f:
            raw = f.read()
        ext = os.path.splitext(path)[1].lower().lstrip('.')
        mime = {'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
                'gif': 'image/gif', 'webp': 'image/webp', 'svg': 'image/svg+xml'}.get(ext, 'image/png')
        return 'data:%s;base64,%s' % (mime, base64.b64encode(raw).decode('ascii')), mime
    except Exception:
        return None, None


# -- HTML ----------------------------------------------------------------------

def build_html(title, sections, lang='fr', embed_images=True):
    """Construit un document HTML autonome et accessible (titre, langue, headings,
    images avec alt). embed_images=True intègre les images en data-URI (utile pour
    coller dans un e-mail)."""
    sections = _norm_sections(sections)
    parts = [
        '<!DOCTYPE html>',
        '<html lang="%s">' % _html.escape(lang),
        '<head><meta charset="utf-8">',
        '<title>%s</title>' % _html.escape(title or 'Document'),
        '</head>',
        '<body>',
        '<h1>%s</h1>' % _html.escape(title or 'Document'),
    ]
    for s in sections:
        if s.get('titre'):
            parts.append('<h2>%s</h2>' % _html.escape(str(s['titre'])))
        for para in _paras(s.get('texte')):
            parts.append('<p>%s</p>' % _html.escape(para).replace('\n', '<br>'))
        if s.get('image'):
            alt = _html.escape(str(s.get('alt') or 'Image'))
            src = None
            if embed_images:
                src, _ = _img_data_uri(s['image'])
            if not src:
                src = _html.escape(str(s['image']))
            parts.append('<figure><img src="%s" alt="%s"><figcaption>%s</figcaption></figure>' % (src, alt, alt))
    parts.append('</body></html>')
    return ('\n'.join(parts)).encode('utf-8'), 'html'


# -- TXT -----------------------------------------------------------------------

def build_txt(title, sections, lang='fr'):
    sections = _norm_sections(sections)
    lines = [title or 'Document', '=' * len(title or 'Document'), '']
    for s in sections:
        if s.get('titre'):
            lines += [str(s['titre']), '-' * len(str(s['titre']))]
        for para in _paras(s.get('texte')):
            lines += [para, '']
        if s.get('image'):
            lines += ['[Image : %s]' % (s.get('alt') or s.get('image')), '']
    return ('\n'.join(lines)).encode('utf-8'), 'txt'


# -- DOCX ----------------------------------------------------------------------

def build_docx(title, sections, lang='fr'):
    from docx import Document
    from docx.shared import Inches
    sections = _norm_sections(sections)
    doc = Document()
    try:
        doc.core_properties.title = title or 'Document'
        doc.core_properties.language = lang
    except Exception:
        pass
    try:
        from docx.oxml.ns import qn
        rpr = doc.styles['Normal'].element.get_or_add_rPr()
        langel = rpr.makeelement(qn('w:lang'), {qn('w:val'): lang})
        rpr.append(langel)
    except Exception:
        pass
    doc.add_heading(title or 'Document', level=0)
    for s in sections:
        if s.get('titre'):
            doc.add_heading(str(s['titre']), level=1)
        for para in _paras(s.get('texte')):
            doc.add_paragraph(para)
        if s.get('image'):
            try:
                doc.add_picture(s['image'], width=Inches(5.5))
            except Exception:
                pass
            cap = doc.add_paragraph()
            run = cap.add_run('Description : ' + str(s.get('alt') or 'Image'))
            run.italic = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), 'docx'


# -- PDF -----------------------------------------------------------------------

def build_pdf(title, sections, lang='fr'):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    sections = _norm_sections(sections)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=(title or 'Document'),
                            leftMargin=2 * cm, rightMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)
    try:
        doc.lang = lang
    except Exception:
        pass
    styles = getSampleStyleSheet()
    story = [Paragraph(_html.escape(title or 'Document'), styles['Title']), Spacer(1, 12)]
    for s in sections:
        if s.get('titre'):
            story.append(Paragraph(_html.escape(str(s['titre'])), styles['Heading1']))
        for para in _paras(s.get('texte')):
            story.append(Paragraph(_html.escape(para).replace('\n', '<br/>'), styles['BodyText']))
            story.append(Spacer(1, 6))
        if s.get('image'):
            try:
                from reportlab.lib.utils import ImageReader
                _iw, _ih = ImageReader(s['image']).getSize()
                _w = min(14 * cm, float(_iw))
                _h = _w * float(_ih) / float(_iw) if _iw else _w
                story.append(RLImage(s['image'], width=_w, height=_h))
            except Exception:
                pass
            story.append(Paragraph('<i>Description : %s</i>' % _html.escape(str(s.get('alt') or 'Image')), styles['BodyText']))
            story.append(Spacer(1, 8))
    doc.build(story)
    return buf.getvalue(), 'pdf'


# -- EPUB ----------------------------------------------------------------------

def build_epub(title, sections, lang='fr'):
    sections = _norm_sections(sections)
    title_e = _html.escape(title or 'Document')
    images = []  # (id, filename, mime, bytes)
    body = ['<h1>%s</h1>' % title_e]
    for s in sections:
        if s.get('titre'):
            body.append('<h2>%s</h2>' % _html.escape(str(s['titre'])))
        for para in _paras(s.get('texte')):
            body.append('<p>%s</p>' % _html.escape(para).replace('\n', '<br/>'))
        if s.get('image'):
            alt = _html.escape(str(s.get('alt') or 'Image'))
            try:
                with open(s['image'], 'rb') as f:
                    raw = f.read()
                ext = os.path.splitext(s['image'])[1].lower().lstrip('.') or 'png'
                mime = {'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
                        'gif': 'image/gif', 'webp': 'image/webp'}.get(ext, 'image/png')
                iid = 'img%d' % len(images)
                fname = '%s.%s' % (iid, 'jpg' if ext == 'jpeg' else ext)
                images.append((iid, fname, mime, raw))
                body.append('<figure><img src="%s" alt="%s"/><figcaption>%s</figcaption></figure>' % (fname, alt, alt))
            except Exception:
                body.append('<p><em>Image : %s</em></p>' % alt)
    content = ('<?xml version="1.0" encoding="UTF-8"?>\n'
               '<!DOCTYPE html>\n'
               '<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="%s" lang="%s">\n'
               '<head><meta charset="UTF-8"/><title>%s</title></head>\n'
               '<body>%s</body></html>' % (lang, lang, title_e, '\n'.join(body)))
    nav = ('<?xml version="1.0" encoding="UTF-8"?>\n'
           '<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="%s">\n'
           '<head><meta charset="UTF-8"/><title>Sommaire</title></head>\n'
           '<body><nav epub:type="toc"><h1>Sommaire</h1><ol><li><a href="content.xhtml">%s</a></li></ol></nav></body>\n'
           '</html>' % (lang, title_e))
    manifest_imgs = ''.join('<item id="%s" href="%s" media-type="%s"/>' % (iid, fn, mime)
                            for (iid, fn, mime, _b) in images)
    opf = ('<?xml version="1.0" encoding="UTF-8"?>\n'
           '<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="bookid">\n'
           '<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">\n'
           '<dc:identifier id="bookid">nimm-doc</dc:identifier>\n'
           '<dc:title>%s</dc:title><dc:language>%s</dc:language>\n'
           '</metadata>\n'
           '<manifest>\n'
           '<item id="content" href="content.xhtml" media-type="application/xhtml+xml"/>\n'
           '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>\n'
           '%s</manifest>\n'
           '<spine><itemref idref="content"/></spine>\n'
           '</package>' % (title_e, lang, manifest_imgs))
    container = ('<?xml version="1.0"?>\n'
                 '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">\n'
                 '<rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>\n'
                 '</container>')
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as z:
        z.writestr(zipfile.ZipInfo('mimetype'), 'application/epub+zip')
        z.writestr('META-INF/container.xml', container.encode('utf-8'))
        z.writestr('OEBPS/content.opf', opf.encode('utf-8'))
        z.writestr('OEBPS/content.xhtml', content.encode('utf-8'))
        z.writestr('OEBPS/nav.xhtml', nav.encode('utf-8'))
        for (iid, fn, mime, raw) in images:
            z.writestr('OEBPS/%s' % fn, raw)
    return buf.getvalue(), 'epub'


# -- Dispatcher ----------------------------------------------------------------

_BUILDERS = {
    'html': build_html, 'txt': build_txt, 'docx': build_docx,
    'pdf': build_pdf, 'epub': build_epub,
}


def build_document(title, sections, fmt='docx', lang='fr'):
    """Construit un document accessible. Retourne (bytes, extension).
    Lève ValueError si le format est inconnu."""
    fmt = (fmt or 'docx').lower().strip()
    if fmt in ('htm',):
        fmt = 'html'
    if fmt not in _BUILDERS:
        raise ValueError("Format inconnu : %s (html, txt, docx, pdf, epub)" % fmt)
    return _BUILDERS[fmt](title or 'Document', sections, lang=lang)
