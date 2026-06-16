# -*- coding: utf-8 -*-
"""
Export de messages NIMM vers différents formats.

Formats supportés : txt, docx, pdf, rtf, odt, epub, mp3.
- txt, rtf, odt, epub : aucune dépendance supplémentaire.
- docx : python-docx (déjà dans requirements.txt).
- pdf  : fpdf2 (pip install fpdf2).
- mp3  : edge-tts (déjà dans requirements.txt).
"""
import datetime
import io
import zipfile

# ─── helpers ──────────────────────────────────────────────────────────────────

def _label(role: str) -> str:
    return "Moi" if role == "user" else "NIMM"


def _now_str() -> str:
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")


def _full_text(items: list) -> str:
    """Texte brut de la conversation (rôle + contenu)."""
    lines = []
    for it in items:
        lines.append(f"[{_label(it['role'])}]")
        lines.append(it["content"].strip())
        lines.append("")
    return "\n".join(lines)


# ─── TXT ──────────────────────────────────────────────────────────────────────

def _to_txt(items: list, ts: str):
    data = _full_text(items).encode("utf-8")
    return data, f"export_nimm_{ts}.txt", "text/plain; charset=utf-8"


# ─── RTF ──────────────────────────────────────────────────────────────────────

def _rtf_escape(text: str) -> str:
    out = []
    for ch in text:
        cp = ord(ch)
        if ch == "\\":
            out.append("\\\\")
        elif ch == "{":
            out.append("\\{")
        elif ch == "}":
            out.append("\\}")
        elif ch == "\n":
            out.append("\\par\n")
        elif cp > 127:
            # Encodage unicode RTF : \uN? (? = substitut ASCII)
            out.append(f"\\u{cp}?")
        else:
            out.append(ch)
    return "".join(out)


def _to_rtf(items: list, ts: str):
    buf = [
        r"{\rtf1\ansi\ansicpg1252\deff0",
        r"{\fonttbl{\f0\froman\fcharset0 Times New Roman;}}",
        r"\widowctrl\wpaper12240\wpaperh15840\margl1800\margr1800\margt1440\margb1440",
        r"\f0\fs22",
    ]
    for it in items:
        role_rtf = _rtf_escape(_label(it["role"]))
        content_rtf = _rtf_escape(it["content"].strip())
        buf.append(r"{\b " + role_rtf + r"}\par")
        buf.append(content_rtf + r"\par\par")
    buf.append("}")
    data = "\n".join(buf).encode("latin-1", errors="replace")
    return data, f"export_nimm_{ts}.rtf", "application/rtf"


# ─── ODT ──────────────────────────────────────────────────────────────────────

def _xml_esc(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _to_odt(items: list, ts: str):
    content_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
    content_xml += '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
    content_xml += 'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
    content_xml += 'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" '
    content_xml += 'office:version="1.3"><office:body><office:text>'

    for it in items:
        role = _xml_esc(_label(it["role"]))
        paragraphs = it["content"].strip().split("\n")
        content_xml += f'<text:p text:style-name="Heading_20_2">{role}</text:p>'
        for para in paragraphs:
            content_xml += f'<text:p text:style-name="Text_20_Body">{_xml_esc(para)}</text:p>'

    content_xml += "</office:text></office:body></office:document-content>"

    styles_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
    styles_xml += '<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
    styles_xml += 'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" '
    styles_xml += 'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" '
    styles_xml += 'office:version="1.3"><office:styles>'
    styles_xml += '<style:style style:name="Text_20_Body" style:family="paragraph"><style:text-properties fo:font-size="11pt"/></style:style>'
    styles_xml += '<style:style style:name="Heading_20_2" style:family="paragraph"><style:text-properties fo:font-weight="bold" fo:font-size="12pt"/></style:style>'
    styles_xml += "</office:styles></office:document-styles>"

    manifest_xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
    manifest_xml += '<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.3">'
    manifest_xml += '<manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/>'
    manifest_xml += '<manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/>'
    manifest_xml += '<manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/>'
    manifest_xml += "</manifest:manifest>"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        # mimetype DOIT être le premier fichier et non compressé
        z.writestr(zipfile.ZipInfo("mimetype"), "application/vnd.oasis.opendocument.text")
        z.writestr("META-INF/manifest.xml", manifest_xml.encode("utf-8"))
        z.writestr("content.xml", content_xml.encode("utf-8"))
        z.writestr("styles.xml", styles_xml.encode("utf-8"))

    return buf.getvalue(), f"export_nimm_{ts}.odt", "application/vnd.oasis.opendocument.text"


# ─── EPUB ─────────────────────────────────────────────────────────────────────

def _to_epub(items: list, ts: str):
    uid = f"nimm-export-{ts}"

    # content.html
    body_html = ""
    for it in items:
        role = _xml_esc(_label(it["role"]))
        paragraphs = it["content"].strip().split("\n")
        body_html += f"<h2>{role}</h2>"
        for para in paragraphs:
            body_html += f"<p>{_xml_esc(para)}</p>"

    content_html = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><meta charset="UTF-8"/><title>Export NIMM</title></head>
<body>{body_html}</body>
</html>"""

    opf = f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="uid">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="uid">{uid}</dc:identifier>
    <dc:title>Export NIMM {ts}</dc:title>
    <dc:language>fr</dc:language>
  </metadata>
  <manifest>
    <item id="content" href="content.xhtml" media-type="application/xhtml+xml"/>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
  </manifest>
  <spine><itemref idref="content"/></spine>
</package>"""

    nav = f"""<?xml version="1.0" encoding="UTF-8"?>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
<head><meta charset="UTF-8"/><title>Table des matières</title></head>
<body><nav epub:type="toc"><ol><li><a href="content.xhtml">Conversation</a></li></ol></nav></body>
</html>"""

    container = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>"""

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(zipfile.ZipInfo("mimetype"), "application/epub+zip")
        z.writestr("META-INF/container.xml", container.encode("utf-8"))
        z.writestr("OEBPS/content.opf", opf.encode("utf-8"))
        z.writestr("OEBPS/content.xhtml", content_html.encode("utf-8"))
        z.writestr("OEBPS/nav.xhtml", nav.encode("utf-8"))

    return buf.getvalue(), f"export_nimm_{ts}.epub", "application/epub+zip"


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _to_docx(items: list, ts: str):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = f"Export NIMM {ts}"

    for it in items:
        role_para = doc.add_paragraph()
        run = role_para.add_run(_label(it["role"]))
        run.bold = True
        run.font.size = Pt(12)

        for line in it["content"].strip().split("\n"):
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)

        doc.add_paragraph()  # ligne vide

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), f"export_nimm_{ts}.docx", \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _to_pdf(items: list, ts: str):
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError(
            "fpdf2 n'est pas installé. Lancez : pip install fpdf2"
        )

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    for it in items:
        # Rôle en gras
        pdf.set_font("Helvetica", style="B", size=12)
        pdf.cell(0, 8, _label(it["role"]), new_x="LMARGIN", new_y="NEXT")
        # Contenu
        pdf.set_font("Helvetica", size=11)
        for line in it["content"].strip().split("\n"):
            pdf.multi_cell(0, 6, line)
        pdf.ln(4)

    data = bytes(pdf.output())
    return data, f"export_nimm_{ts}.pdf", "application/pdf"


# ─── MP3 ──────────────────────────────────────────────────────────────────────

async def _to_mp3(items: list, ts: str):
    try:
        import edge_tts
    except ImportError:
        raise RuntimeError(
            "edge-tts n'est pas installé. Lancez : pip install edge-tts"
        )

    text = _full_text(items)
    # Voix française de haute qualité disponible avec edge-tts
    voice = "fr-FR-DeniseNeural"
    communicate = edge_tts.Communicate(text, voice)

    buf = io.BytesIO()
    async for chunk in communicate.stream():
        if chunk["type"] == "audio":
            buf.write(chunk["data"])

    return buf.getvalue(), f"export_nimm_{ts}.mp3", "audio/mpeg"


# ─── Point d'entrée ───────────────────────────────────────────────────────────

async def export_messages(items: list, fmt: str) -> tuple:
    """
    items : list de {role: 'user'|'assistant', content: str}
    fmt   : 'txt' | 'docx' | 'pdf' | 'rtf' | 'odt' | 'epub' | 'mp3'

    Retourne (bytes, filename, mime_type).
    """
    if not items:
        raise ValueError("Aucun message à exporter.")

    ts = _now_str()

    if fmt == "txt":
        return _to_txt(items, ts)
    elif fmt == "docx":
        return _to_docx(items, ts)
    elif fmt == "pdf":
        return _to_pdf(items, ts)
    elif fmt == "rtf":
        return _to_rtf(items, ts)
    elif fmt == "odt":
        return _to_odt(items, ts)
    elif fmt == "epub":
        return _to_epub(items, ts)
    elif fmt == "mp3":
        return await _to_mp3(items, ts)
    else:
        raise ValueError(f"Format non supporté : {fmt}")
