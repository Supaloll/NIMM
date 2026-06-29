# -*- coding: utf-8 -*-
"""
CoaNIMM — agent d'exécution pour NIMM.

Exécute des scripts Python dans un répertoire de travail dédié et sandboxé
(data/coanimm_workspace/), sous contrôle de permissions explicites — voir
core.database.agent_permission_granted :

  - 'once'    : exécute une seule fois sans rien enregistrer ;
  - 'project' : autorise pour le fil de conversation courant (thread_id) ;
  - 'always'  : autorise pour toujours, tous fils confondus.

Si aucune permission n'est accordée et qu'aucun confirm_scope n'est fourni, les
fonctions ci-dessous renvoient {'status': 'permission_required', ...} : c'est au
frontend de demander à l'utilisateur, puis de rappeler avec confirm_scope.

Flux Plan→Explore→Execute :
  1. generate_plan()     : LLM décrit le plan + indique si exploration disque nécessaire
  2. explore_directory() : script lecture-seule pour informer le plan (si needs_explore)
  3. run_generated()     : génère et exécute le script final
"""
import os
import re
import subprocess
import sys
import tempfile

import core.database as db

WORKSPACE_DIRNAME = 'coanimm_workspace'
TIMEOUT_SECONDS = 180

# Extensions reconnues pour le routage automatique des fichiers générés
_IMAGE_EXTS  = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}
_TEXT_EXTS   = {'.txt', '.md', '.csv', '.json', '.html', '.xml', '.log'}
_MAX_TEXT_INLINE = 4000   # chars max injectés dans le résultat outil

GENERATED_ACTION = 'exec_generated_code'
EXPLORE_ACTION   = 'explorer_disque'

PLANNING_SYSTEM_PROMPT = (
    "Tu es CoaNIMM, l'agent d'exécution de NIMM.\n"
    "L'utilisateur va te confier une tâche à automatiser. "
    "Ta réponse doit être UN TEXTE BRUT UNIQUEMENT, lisible tel quel par un utilisateur aveugle "
    "sur une plage braille. Aucune mise en forme, aucune balise, aucun symbole de formatage.\n\n"
    "INTERDICTIONS ABSOLUES dans ta réponse :\n"
    "- Pas d'astérisques, pas de gras (**texte**), pas d'italique (*texte*)\n"
    "- Pas de titres markdown (## ou ###)\n"
    "- Pas de backticks, pas de blocs de code (``` ou `code`)\n"
    "- Pas de tirets de liste (- item) ni de puces\n"
    "- Pas de PowerShell, pas de Python, pas de commandes\n"
    "- Pas de HTML\n\n"
    "FORMAT OBLIGATOIRE de ta réponse (deux blocs séparés par une ligne vide) :\n\n"
    "Bloc 1 — une seule ligne, exactement :\n"
    "  EXPLORER: oui   si la tâche nécessite de lire des dossiers ou fichiers sur le disque.\n"
    "  EXPLORER: non   si la tâche peut être planifiée sans accès disque.\n\n"
    "Bloc 2 — le plan en texte brut, 3 à 8 phrases numérotées (1. 2. 3. ...), en français. "
    "Chaque phrase décrit une étape concrète. Pas de sous-points. Pas de récapitulatif final. "
    "Arrête-toi dès que le plan est complet."
)

EXPLORE_SYSTEM_PROMPT = (
    "Tu es CoaNIMM en mode exploration (lecture seule).\n"
    "Génère un script Python qui explore le système de fichiers et affiche ce qu'il trouve. "
    "LECTURE SEULE UNIQUEMENT — les instructions suivantes sont absolument interdites : "
    "shutil.move, shutil.copy, shutil.copytree, shutil.rmtree, os.rename, os.replace, "
    "os.remove, os.unlink, os.makedirs, os.mkdir, open(..., 'w'), open(..., 'a'), "
    "open(..., 'wb'), open(..., 'ab').\n"
    "Règles habituelles : pas de triple-guillemets, pas de input(), pas de sys.stdin. "
    "Affiche un rapport clair et lisible avec print()."
)

GENERATE_SYSTEM_PROMPT = (
    "Tu es CoaNIMM, l'agent d'exécution de NIMM, un assistant personnel local-first.\n"
    "A partir d'une consigne en langage naturel, génère un script Python autonome qui "
    "réalise cette tâche.\n"
    "Réponds UNIQUEMENT avec le code Python, sans balises markdown (```), sans "
    "explication avant ou après.\n"
    "RÈGLES IMPÉRATIVES :\n"
    "1. Écris un script Python autonome, complet et directement exécutable. "
    "Les triple-guillemets (\"\"\" ou ''') sont autorisés, pour les docstrings comme "
    "pour les chaînes multilignes.\n"
    "2. N'utilise JAMAIS input() ni sys.stdin : le script s'exécute sans terminal interactif. Si tu as besoin d'une réponse de l'utilisateur, utilise le protocole __NIMM_DEMANDE__ (règle 3).\n"
    "3. INTERACTION UTILISATEUR : si la tâche nécessite une validation ou un choix (ex : confirmer un plan avant des opérations irréversibles), le script doit :\n"
    "   a) Afficher le plan/analyse complet avec print()\n"
    "   b) Terminer par exactement ces deux lignes :\n"
    "      print('__NIMM_DEMANDE__: ta question ici')\n"
    "      import sys; sys.exit(0)\n"
    "   CoaNIMM détectera ce marqueur, montrera un champ de saisie à l'utilisateur, et relancera génération + exécution avec la réponse et tout le contexte précédent. Pour les tâches sans risque, exécute directement sans demander.\n"
    "4. Affiche chaque action avec print() au fur et à mesure (ex : 'Déplacé : ancien -> nouveau').\n"
    "Le script s'exécute dans un processus isolé.\n"
    "CONFINEMENT : écris tous tes fichiers de sortie dans le RÉPERTOIRE COURANT. "
    "Toute écriture, suppression ou déplacement hors du répertoire courant est bloqué "
    "par sécurité, sauf si l'utilisateur a explicitement autorisé le dossier visé.\n"
    "BIBLIOTHÈQUES DISPONIBLES : bibliothèque standard Python + reportlab (PDF), "
    "python-docx (Word), PIL/Pillow (images), pandas, openpyxl.\n"
    "ACCESSIBILITÉ (l'utilisateur est aveugle et lit avec un lecteur d'écran et une plage braille) :\n"
    "  - Pour un document destiné à être LU par un lecteur d'écran, privilégie un .docx "
    "(python-docx) avec des styles de titres (add_heading) : c'est nativement accessible.\n"
    "      from docx import Document\n"
    "      d = Document(); d.add_heading('Mon titre', level=1)\n"
    "      d.add_paragraph('Texte du paragraphe.'); d.save('rapport.docx')\n"
    "  - Si un PDF est explicitement demandé, utilise reportlab en mode STRUCTURÉ (platypus) "
    "pour un ordre de lecture correct, avec titre du document, langue française, et texte "
    "alternatif sur les images. Évite canvas.drawString (texte non structuré, illisible au "
    "lecteur d'écran).\n"
    "      from reportlab.lib.pagesizes import A4\n"
    "      from reportlab.lib.styles import getSampleStyleSheet\n"
    "      from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image\n"
    "      doc = SimpleDocTemplate('rapport.pdf', pagesize=A4, title='Mon titre', lang='fr-FR')\n"
    "      s = getSampleStyleSheet()\n"
    "      story = [Paragraph('Mon titre', s['Title']), Spacer(1, 12),\n"
    "               Paragraph('Texte du paragraphe.', s['BodyText'])]\n"
    "      img_path = nimm_generate_image('une illustration de ...')\n"
    "      story += [Image(img_path, width=200, height=150),\n"
    "                Paragraph(\"Description de l'image pour le lecteur d'écran.\", s['BodyText'])]\n"
    "      doc.build(story)\n"
    "HELPER INJECTÉ (disponible sans import) :\n"
    "  nimm_generate_image(prompt: str) -> str\n"
    "  Génère une image IA à partir du prompt et retourne le chemin absolu du fichier PNG "
    "dans le répertoire de travail courant.\n"
    "  nimm_web_search(query: str) -> str\n"
    "  Recherche web : passe une REQUÊTE en langage naturel (jamais une URL), retourne un "
    "texte de résultats. Pour une info à jour ou un exemple.\n"
    "  nimm_github_search(query: str) -> str\n"
    "  Recherche GitHub (dépôts ou code) à partir d'une requête, retourne un texte de "
    "résultats avec liens, pour s'inspirer d'exemples de code.\n"
    "  nimm_search_documents(query: str) -> str\n"
    "  Interroge la BASE DE CONNAISSANCES de l'utilisateur (documents déjà ingérés) et "
    "retourne les passages pertinents. Pour répondre à partir de ses documents.\n"
    "  nimm_extract_text(path: str) -> str\n"
    "  Extrait le texte d'un fichier (PDF, Word, ODT, RTF, EPUB, HTML, image avec OCR). "
    "Lecture seule. Pour résumer ou traiter le contenu d'un document.\n"
    "  nimm_ask_llm(prompt: str, system: str = '') -> str\n"
    "  Demande au LLM une sous-tâche (résumer, classer, traduire, reformuler) et retourne "
    "sa réponse texte. Utile pour traiter du contenu au fil de l'exécution.\n"
    "  nimm_read_url(url: str) -> str\n"
    "  Extrait le texte principal d'une page web précise (protégé anti-SSRF). À distinguer "
    "de nimm_web_search qui, lui, prend une requête.\n"
    "  nimm_translate(text: str, target_lang: str = 'anglais') -> str\n"
    "  Traduit un texte dans la langue cible et retourne la traduction.\n"
    "  nimm_expurgate(text: str, consigne: str = '') -> str\n"
    "  Produit une version ADAPTÉE AUX ENFANTS d'un texte : retire/adoucit les scènes "
    "violentes, sexuelles, d'horreur ou le langage grossier, en préservant l'histoire ; "
    "peut abréger. Retourne le texte adapté.\n"
    "  nimm_coloring_page(subject: str) -> str\n"
    "  Génère un COLORIAGE (dessin au trait noir et blanc, pour enfants) sur le sujet donné "
    "et retourne le chemin du PNG.\n"
    "  nimm_make_document(title: str, sections: list, fmt: str = 'docx', lang: str = 'fr') -> str\n"
    "  Crée un DOCUMENT ACCESSIBLE (titres, langue, images décrites) et retourne son chemin. "
    "fmt : docx, pdf, epub, pptx, html, txt. sections : liste de dicts {'titre':..., 'texte':..., "
    "'image': chemin, 'alt': description}. Utilise html pour un contenu à coller dans un e-mail.\n"
    "  nimm_transcribe(audio_path: str) -> str\n"
    "  Transcrit un fichier audio (voix → texte) via Whisper local et retourne le texte.\n"
    "  nimm_speak(text: str, voice: str = '') -> str\n"
    "  Synthétise un texte en AUDIO (TTS) et retourne le chemin du fichier son. Pour un livre audio.\n"
    "  nimm_describe_image(path: str, prompt: str = '') -> str\n"
    "  Décrit une image (texte alternatif accessible) via le modèle de vision et retourne le texte.\n"
    "  nimm_simplify(text: str, niveau: str = '') -> str\n"
    "  Réécrit un texte en FALC (Facile À Lire et à Comprendre) : phrases courtes, mots simples. Accessibilité cognitive.\n"
    "  nimm_resize_image(path: str, max_width: int = 1200, fmt: str = '') -> str\n"
    "  Redimensionne et/ou convertit une image (jpg, png, webp…) et retourne le chemin du fichier produit.\n"
    "  nimm_anonymize(text: str) -> str\n"
    "  Masque les données personnelles d'un texte (noms, e-mails, téléphones, adresses…) et retourne le texte anonymisé.\n"
    "  nimm_merge_pdf(paths: list, name: str = '') -> str\n"
    "  Fusionne plusieurs fichiers PDF (liste de chemins) en un seul et retourne le chemin du PDF produit.\n"
    "  nimm_split_pdf(path: str, pages: str) -> str\n"
    "  Extrait des pages d'un PDF (ex. pages='1-3,5') dans un nouveau PDF et retourne son chemin.\n"
    "  nimm_pdf_from_images(paths: list, name: str = '') -> str\n"
    "  Assemble une liste d'images en un PDF (une image par page) et retourne le chemin du PDF.\n"
    "  nimm_read_table(path: str) -> str\n"
    "  Lit un fichier CSV/TSV et le renvoie en tableau Markdown lisible.\n"
    "  nimm_ocr_document(path: str = '', url: str = '') -> str\n"
    "  Extrait le texte d'un PDF ou d'une image via Mistral OCR (mistral-ocr-latest) en préservant la structure (titres, tableaux, formules). Retourne le texte en Markdown.\n"
    "  nimm_mistral_speak(text: str, voice_id: str = '', ref_audio_path: str = '') -> str\n"
    "  Synthèse vocale Mistral : voix préréglée (voice_id) ou clonage zero-shot (ref_audio_path = chemin d'un fichier audio). Retourne le chemin du fichier audio produit.\n"
    "  nimm_audio_overview(content: str, voice1: str = '', voice2: str = '') -> str\n"
    "  Crée un RÉSUMÉ AUDIO façon podcast : génère un dialogue à 2 voix sur le contenu puis le synthétise (Gemini TTS). Retourne le chemin du fichier audio.\n"
    "  nimm_list_voices() -> list\n"
    "  Retourne la liste des voix TTS disponibles (Edge, Gemini, Voxtral…) avec leur id et label. "
    "Utiliser avant nimm_make_daisy ou nimm_speak pour choisir la bonne voix.\n"
    "  nimm_make_daisy(title: str, sections: list, lang: str = 'fr', voice: str = '', style: str = '') -> str\n"
    "  Crée un LIVRE AUDIO DAISY 2.02 (format accessible standard Victor Reader, AMIS, EasyReader). "
    "Produit un fichier .daisy (ZIP) avec ncc.html, fichiers SMIL et audio MP3 synchronisés. "
    "sections : liste de dicts {'titre': str, 'texte': str}. "
    "voice : id de voix (appeler nimm_list_voices() pour les options y compris voix Voxtral personnalisées) ; "
    "vide = voix par défaut. "
    "Retourne le chemin du fichier .daisy.\n"
    "N'importe aucun de ces helpers (nimm_generate_image, nimm_web_search, nimm_github_search, "
    "nimm_search_documents, nimm_extract_text, nimm_ask_llm, nimm_read_url, nimm_translate, "
    "nimm_expurgate, nimm_coloring_page, nimm_make_document, nimm_transcribe, nimm_speak, "
    "nimm_describe_image, nimm_simplify, nimm_resize_image, nimm_anonymize, nimm_merge_pdf, "
    "nimm_split_pdf, nimm_pdf_from_images, nimm_read_table, nimm_audio_overview, "
    "nimm_make_daisy, nimm_list_voices, nimm_ocr_document, nimm_mistral_speak) : "
    "ils sont déjà présents dans l'environnement."
)

SKILL_WRITER_SYSTEM_PROMPT = (
    "Tu es CoaNIMM en mode rédaction de fiche skill.\n"
    "Une fiche skill capture une MÉTHODE qui vient d'être validée par Laurent, pour pouvoir "
    "la redemander plus tard. Ce n'est pas un script figé : c'est un mode d'emploi en langage "
    "naturel qui enseigne la LOGIQUE de la méthode.\n\n"
    "RÈGLE CARDINALE : enseigne par la logique de la méthode, jamais en recopiant l'exemple. "
    "Une fiche collée au cas précis (« convertir CETTE image-ci ») ne sert à rien la fois "
    "suivante ; une fiche trop vague (« retoucher des images ») perd les nuances. Capture le "
    "PRINCIPE qui a fait que ça a marché (par exemple « seuillage binaire » ou « quantification "
    "de palette »), pas les valeurs précises du test ni le fichier d'exemple particulier.\n\n"
    "Tu reçois en entrée la consigne d'origine de Laurent et le script validé. Produis la fiche.\n\n"
    "FORMAT OBLIGATOIRE de ta réponse, en TEXTE BRUT uniquement (Laurent lit sur plage braille : "
    "aucun markdown, aucune balise, aucun astérisque, aucun titre #, aucun backtick, aucune puce) :\n\n"
    "Ligne 1, exactement, la description « quand l'utiliser » en une seule phrase :\n"
    "  DESCRIPTION: <une phrase qui dit dans quel cas ce skill s'applique>\n"
    "Ligne 2, exactement, les déclencheurs séparés par des virgules :\n"
    "  MOTS-CLES: <mot1, mot2, mot3, ...>\n"
    "Puis une ligne vide, puis le CORPS de la fiche : la méthode en langage naturel, "
    "3 à 8 phrases, qui explique le principe et comment l'appliquer. Ne mentionne pas le "
    "fichier d'exemple. Ne recopie pas le code.\n\n"
    "Réponds UNIQUEMENT « SKIP » (et rien d'autre) si l'entrée ne décrit pas une méthode "
    "réutilisable (résultat improvisé une seule fois, sans principe généralisable)."
)


def _sanitize_dirname(name: str) -> str:
    """Nettoie un nom de fil pour en faire un nom de dossier valide (Windows compris)."""
    name = (name or '').strip()
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', name)
    name = name.strip(' .')
    return name[:60] or 'sans_titre'


def _workspace_dir(thread_id: str = None) -> str:
    """Répertoire de travail UNIQUE de CoaNIMM — atelier global, indépendant du fil
    de conversation. Le thread_id est ignoré (un seul espace de travail partagé) :
    CoaNIMM est une surface autonome, pas une notion par-fil."""
    base = os.path.join(db.DATA_DIR, WORKSPACE_DIRNAME)
    os.makedirs(base, exist_ok=True)
    return base


def purge_workspace() -> dict:
    """Vide l'espace de travail global de CoaNIMM (fichiers produits + sous-dossiers),
    en conservant le dossier lui-même. Action explicite, utile après une session
    confidentielle. Retourne {'status':'ok', 'removed': n}."""
    import shutil
    base = _workspace_dir()
    removed = 0
    try:
        for name in os.listdir(base):
            path = os.path.join(base, name)
            try:
                if os.path.isdir(path) and not os.path.islink(path):
                    shutil.rmtree(path)
                else:
                    os.remove(path)
                removed += 1
            except Exception as _e:
                print(f"[COANIMM] purge : impossible de supprimer {path} : {_e}")
    except Exception as _e:
        return {'status': 'error', 'message': str(_e)}
    print(f"[COANIMM] Espace de travail purgé : {removed} élément(s).")
    return {'status': 'ok', 'removed': removed}


def _strip_code_fences(text: str) -> str:
    """Extrait le code Python d'une réponse LLM, même imparfaite.

    Gère trois cas fréquents qui faisaient échouer l'exécution :
      - balises ```python ... ``` situées n'importe où (et pas seulement en tête) ;
      - texte explicatif avant ou après le bloc, malgré la consigne ;
      - plusieurs blocs : on retient le plus long (le script complet) ;
      - réponse tronquée par max_tokens : un ``` d'ouverture sans fermeture est
        nettoyé pour récupérer le code partiel (qui déclenchera ensuite le retry).
    """
    text = (text or '').strip()
    if not text:
        return ''
    # Blocs ```lang\n ... ``` complets, où qu'ils soient
    blocks = re.findall(r'```[a-zA-Z0-9_+\-]*\n?(.*?)```', text, re.DOTALL)
    if blocks:
        return max(blocks, key=len).strip()
    # Pas de bloc fermé : retirer d'éventuelles lignes ``` orphelines (tête/queue)
    lines = text.splitlines()
    if lines and lines[0].lstrip().startswith('```'):
        lines = lines[1:]
    if lines and lines[-1].rstrip().endswith('```'):
        lines = lines[:-1]
    return '\n'.join(lines).strip()


def _check_syntax(code: str):
    """Vérifie la syntaxe du code avant exécution. Retourne un message d'erreur ou None."""
    try:
        compile(code, '<generated>', 'exec')
        return None
    except SyntaxError as e:
        return f"Erreur de syntaxe ligne {e.lineno} : {e.msg}"


def _analyze_code_risks(code: str) -> list:
    """Analyse statique AST du code généré, pour AFFICHER les avertissements dans l'UI.

    Délègue à modules.coanimm_safety.risks_for_display afin de partager UNE SEULE
    source de vérité avec le blocage/confirmation (classify_for_execution) : les
    avertissements montrés à l'utilisateur correspondent ainsi exactement à ce qui
    est réellement bloqué, à confirmer, ou confiné. Format : [{'level','message'}].
    """
    try:
        import modules.coanimm_safety as _safety
        return _safety.risks_for_display(code)
    except Exception:
        return []


def _build_prologue(thread_id: str, workdir: str) -> str:
    """Construit le code Python injecte en tete de chaque script CoaNIMM.

    Definit les helpers d'outils (image, recherche web, GitHub) appelant les endpoints
    locaux. Un outil DESACTIVE dans les reglages est remplace par un stub qui leve une
    erreur claire (au lieu d'etre silencieusement absent)."""
    tid = (thread_id or "").replace("'", "")
    try:
        _disabled = set(db.list_coanimm_disabled_tools())
    except Exception:
        _disabled = set()
    header = "import urllib.request as _nimm_ur, json as _nimm_json\n"
    img = (
        "def nimm_generate_image(prompt, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"prompt\": prompt, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/generate_image\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_generate_image : \" + _res.get(\"message\", \"?\"))\n"
        "    print(\"Image générée : \" + _res[\"filepath\"])\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    web = (
        "def nimm_web_search(query, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"query\": query, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/web_search\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=60) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    gh = (
        "def nimm_github_search(query, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"query\": query, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/github_search\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=60) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    def _stub(fn, label):
        return ("def %s(*_a, **_k):\n"
                "    raise RuntimeError(\"Outil désactivé dans les réglages CoaNIMM : %s.\")\n") % (fn, label)
    parts = [header]
    parts.append(img if "image" not in _disabled else _stub("nimm_generate_image", "génération d'image"))
    parts.append(web if "web" not in _disabled else _stub("nimm_web_search", "recherche web"))
    parts.append(gh if "github" not in _disabled else _stub("nimm_github_search", "recherche GitHub"))
    ds = (
        "def nimm_search_documents(query, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"query\": query, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/doc_search\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=60) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    ex = (
        "def nimm_extract_text(path, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"path\": path, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/extract_text\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=180) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    al = (
        "def nimm_ask_llm(prompt, system='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"prompt\": prompt, \"system\": system, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/ask_llm\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=180) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    ru = (
        "def nimm_read_url(url, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"url\": url, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/read_url\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(ds if "doc_search" not in _disabled else _stub("nimm_search_documents", "consulter la base de connaissances"))
    parts.append(ex if "extract_text" not in _disabled else _stub("nimm_extract_text", "extraire le texte d'un document"))
    parts.append(al if "ask_llm" not in _disabled else _stub("nimm_ask_llm", "sous-tache IA"))
    parts.append(ru if "read_url" not in _disabled else _stub("nimm_read_url", "lire une page web"))
    tr = (
        "def nimm_translate(text, target_lang='anglais', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"text\": text, \"target_lang\": target_lang, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/translate\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    exp = (
        "def nimm_expurgate(text, consigne='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"text\": text, \"consigne\": consigne, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/expurgate\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=180) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    col = (
        "def nimm_coloring_page(subject, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"subject\": subject, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/coloring_page\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_coloring_page : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    parts.append(tr if "translate" not in _disabled else _stub("nimm_translate", "traduire"))
    parts.append(exp if "expurgate" not in _disabled else _stub("nimm_expurgate", "expurger un texte"))
    parts.append(col if "coloring" not in _disabled else _stub("nimm_coloring_page", "coloriage"))
    md = (
        "def nimm_make_document(title, sections, fmt='docx', lang='fr', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"title\": title, \"sections\": sections, \"fmt\": fmt, \"lang\": lang, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/make_document\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=180) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_make_document : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    parts.append(md if "make_document" not in _disabled else _stub("nimm_make_document", "creer un document"))
    tx = (
        "def nimm_transcribe(audio_path, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"path\": audio_path, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/transcribe\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=600) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(tx if "transcribe" not in _disabled else _stub("nimm_transcribe", "transcrire un audio"))
    sp = (
        "def nimm_list_voices(_tid='%s'):\n"
        "    import requests\n"
        "    _res = requests.get('http://localhost:8080/api/coanimm/list_voices', timeout=10).json()\n"
        "    return _res.get('voices', [])\n"
        % tid,
        "def nimm_speak(text, voice='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"text\": text, \"voice\": voice, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/speak\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=300) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_speak : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    di = (
        "def nimm_describe_image(path, prompt='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"path\": path, \"prompt\": prompt, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/describe_image\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(sp if "speak" not in _disabled else _stub("nimm_speak", "donner la voix"))
    parts.append(di if "describe_image" not in _disabled else _stub("nimm_describe_image", "decrire une image"))
    si = (
        "def nimm_simplify(text, niveau='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"text\": text, \"niveau\": niveau, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/simplify\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=180) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    ri = (
        "def nimm_resize_image(path, max_width=1200, fmt='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"path\": path, \"max_width\": max_width, \"fmt\": fmt, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/resize_image\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_resize_image : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    parts.append(si if "simplify" not in _disabled else _stub("nimm_simplify", "simplifier un texte"))
    parts.append(ri if "resize_image" not in _disabled else _stub("nimm_resize_image", "redimensionner une image"))
    an = (
        "def nimm_anonymize(text, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"text\": text, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/anonymize\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=180) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    mp = (
        "def nimm_merge_pdf(paths, name='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"paths\": paths, \"name\": name, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/merge_pdf\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_merge_pdf : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    parts.append(an if "anonymize" not in _disabled else _stub("nimm_anonymize", "anonymiser un texte"))
    parts.append(mp if "merge_pdf" not in _disabled else _stub("nimm_merge_pdf", "fusionner des PDF"))
    sd = (
        "def nimm_split_pdf(path, pages, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"path\": path, \"pages\": pages, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/split_pdf\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_split_pdf : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    pi = (
        "def nimm_pdf_from_images(paths, name='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"paths\": paths, \"name\": name, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/pdf_from_images\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_pdf_from_images : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    parts.append(sd if "split_pdf" not in _disabled else _stub("nimm_split_pdf", "decouper un PDF"))
    parts.append(pi if "pdf_from_images" not in _disabled else _stub("nimm_pdf_from_images", "creer un PDF depuis des images"))
    rt = (
        "def nimm_read_table(path, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"path\": path, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/read_table\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=120) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(rt if "read_table" not in _disabled else _stub("nimm_read_table", "lire un tableau"))
    ed = (
        "def nimm_expurgate_doc(path, consigne='', fmt='docx', allow_cloud=False, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"path\": path, \"consigne\": consigne, \"fmt\": fmt, \"allow_cloud\": allow_cloud, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/expurgate_document\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=300) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(ed if "expurgate_doc" not in _disabled else _stub("nimm_expurgate_doc", "expurger un document"))
    fim = (
        "def nimm_codestral_fim(prefix, suffix='', stop=None, temperature=0.0, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"prefix\": prefix, \"suffix\": suffix, \"stop\": stop or [], \"temperature\": temperature, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/codestral_fim\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=60) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(fim if "codestral_fim" not in _disabled else _stub("nimm_codestral_fim", "completer du code (FIM)"))
    ocr = (
        "def nimm_ocr_document(path='', url='', _tid='%s'):\n"
        "    import urllib.request as _ur_ocr, json as _j_ocr, os as _os_ocr, base64 as _b64_ocr\n"
        "    if not path and not url:\n"
        "        raise ValueError('nimm_ocr_document: fournir path ou url')\n"
        "    if path:\n"
        "        _raw = open(path, 'rb').read()\n"
        "        _ext = _os_ocr.path.splitext(path)[1].lower().lstrip('.')\n"
        "        _mime = 'application/pdf' if _ext == 'pdf' else f'image/{_ext or \"octet-stream\"}'\n"
        "        _bname = _os_ocr.path.basename(path)\n"
        "        _b64s = _b64_ocr.b64encode(_raw).decode()\n"
        "        boundary = b'----NimmOcrBnd'\n"
        "        _body = (b'--' + boundary + b'\r\n'\n"
        "            + f'Content-Disposition: form-data; name=\"file\"; filename=\"{_bname}\"\r\n'.encode()\n"
        "            + f'Content-Type: {_mime}\r\n\r\n'.encode()\n"
        "            + _raw + b'\r\n' + b'--' + boundary + b'--\r\n')\n"
        "        _req_ocr = _ur_ocr.Request(\n"
        "            'http://localhost:8080/api/mistral/ocr',\n"
        "            data=_body,\n"
        "            headers={'Content-Type': f'multipart/form-data; boundary={boundary.decode()}'},\n"
        "        )\n"
        "        with _ur_ocr.urlopen(_req_ocr, timeout=120) as _rr:\n"
        "            return _j_ocr.loads(_rr.read()).get('text', '')\n"
        "    else:\n"
        "        _d2 = f'url={_ur_ocr.quote(url)}&thread_id={_tid}'.encode()\n"
        "        _rq2 = _ur_ocr.Request('http://localhost:8080/api/mistral/ocr',\n"
        "            data=_d2, headers={'Content-Type': 'application/x-www-form-urlencoded'})\n"
        "        with _ur_ocr.urlopen(_rq2, timeout=120) as _rr2:\n"
        "            return _j_ocr.loads(_rr2.read()).get('text', '')\n"
    ) % tid
    parts.append(ocr if "ocr_document" not in _disabled else _stub("nimm_ocr_document", "OCR Mistral (PDF/image)"))
    mspk = (
        "def nimm_mistral_speak(text, voice_id='', ref_audio_path='', _tid='%s'):\n"
        "    import urllib.request as _ur_ms, urllib.parse as _up_ms, json as _j_ms, os as _os_ms\n"
        "    _fields = [('text', text), ('voice_id', voice_id), ('fmt', 'mp3'), ('thread_id', _tid)]\n"
        "    if ref_audio_path and _os_ms.path.isfile(ref_audio_path):\n"
        "        import base64 as _b64_ms\n"
        "        _ref = _b64_ms.b64encode(open(ref_audio_path, 'rb').read()).decode()\n"
        "        boundary = b'----NimmMsSpeakBnd'\n"
        "        _bname = _os_ms.path.basename(ref_audio_path)\n"
        "        _ext_ms = _os_ms.path.splitext(ref_audio_path)[1].lower().lstrip('.') or 'wav'\n"
        "        _mime_ms = f'audio/{_ext_ms}'\n"
        "        _raw_ref = open(ref_audio_path, 'rb').read()\n"
        "        _body_ms = (b'--' + boundary + b'\r\n'\n"
        "            + f'Content-Disposition: form-data; name=\"text\"\r\n\r\n'.encode() + text.encode() + b'\r\n'\n"
        "            + b'--' + boundary + b'\r\n'\n"
        "            + f'Content-Disposition: form-data; name=\"ref_audio\"; filename=\"{_bname}\"\r\n'.encode()\n"
        "            + f'Content-Type: {_mime_ms}\r\n\r\n'.encode() + _raw_ref + b'\r\n'\n"
        "            + b'--' + boundary + b'--\r\n')\n"
        "        _req_ms = _ur_ms.Request('http://localhost:8080/api/mistral/audio/speak',\n"
        "            data=_body_ms, headers={'Content-Type': f'multipart/form-data; boundary={boundary.decode()}'])\n"
        "    else:\n"
        "        _data_ms = _up_ms.urlencode(dict(_fields)).encode()\n"
        "        _req_ms = _ur_ms.Request('http://localhost:8080/api/mistral/audio/speak',\n"
        "            data=_data_ms, headers={'Content-Type': 'application/x-www-form-urlencoded'})\n"
        "    with _ur_ms.urlopen(_req_ms, timeout=60) as _rms:\n"
        "        _audio_ms = _rms.read()\n"
        "    import tempfile as _tf_ms\n"
        "    _out_ms = _tf_ms.NamedTemporaryFile(suffix='.mp3', delete=False)\n"
        "    _out_ms.write(_audio_ms); _out_ms.close()\n"
        "    return _out_ms.name\n"
    ) % tid
    parts.append(mspk if "mistral_speak" not in _disabled else _stub("nimm_mistral_speak", "synthese vocale Mistral"))
    ao = (
        "def nimm_audio_overview(content, voice1='', voice2='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"content\": content, \"voice1\": voice1, \"voice2\": voice2, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/audio_overview\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=300) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_audio_overview : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    parts.append(ao if "audio_overview" not in _disabled else _stub("nimm_audio_overview", "resume audio"))
    dy = (
        "def nimm_make_daisy(title, sections, lang='fr', voice='', style='', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"title\": title, \"sections\": sections, \"lang\": lang, \"voice\": voice, \"style\": style, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/make_daisy\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=600) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_make_daisy : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    parts.append(dy if "make_daisy" not in _disabled else _stub("nimm_make_daisy", "creer un livre DAISY"))
    qr = (
        "def nimm_qr_code(content, qr_type='text', name='', vcard_phone='', vcard_email='',"
        " vcard_org='', vcard_url='', vcard_note='', wifi_ssid='', wifi_password='', wifi_security='WPA', _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"content\": content, \"qr_type\": qr_type, \"name\": name,"
        " \"vcard_phone\": vcard_phone, \"vcard_email\": vcard_email, \"vcard_org\": vcard_org,"
        " \"vcard_url\": vcard_url, \"vcard_note\": vcard_note, \"wifi_ssid\": wifi_ssid,"
        " \"wifi_password\": wifi_password, \"wifi_security\": wifi_security,"
        " \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/qr_code\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=30) as _r:\n"
        "        _res = _nimm_json.loads(_r.read())\n"
        "    if _res.get(\"status\") != \"ok\":\n"
        "        raise RuntimeError(\"nimm_qr_code : \" + _res.get(\"message\", \"?\"))\n"
        "    return _res[\"filepath\"]\n"
    ) % tid
    wp = (
        "def nimm_wikipedia(query, lang='fr', sentences=5, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"query\": query, \"lang\": lang, \"sentences\": sentences,"
        " \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/wikipedia\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=30) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    wd = (
        "def nimm_wikidata(query, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"query\": query, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/wikidata\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=30) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(qr if "qr_code" not in _disabled else _stub("nimm_qr_code", "generer un QR code"))
    parts.append(wp if "wikipedia" not in _disabled else _stub("nimm_wikipedia", "rechercher sur Wikipedia"))
    parts.append(wd if "wikidata" not in _disabled else _stub("nimm_wikidata", "interroger Wikidata"))
    sr = (
        "def nimm_sirene(query, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"query\": query, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/sirene\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=15) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    dg = (
        "def nimm_datagouv(query, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"query\": query, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/datagouv\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=15) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    mt = (
        "def nimm_meteo(location, days=3, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"location\": location, \"days\": days, \"thread_id\": _tid}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/meteo\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=15) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(sr if "sirene" not in _disabled else _stub("nimm_sirene", "rechercher une entreprise INSEE"))
    parts.append(dg if "datagouv" not in _disabled else _stub("nimm_datagouv", "rechercher sur data.gouv.fr"))
    parts.append(mt if "meteo" not in _disabled else _stub("nimm_meteo", "obtenir la meteo"))

    ma = (
        "def nimm_mistral_agent(message, agent_id, conversation_id=None, _tid='%s'):\n"
        "    _data = _nimm_json.dumps({\"agent_id\": agent_id, \"message\": message,\n"
        "        \"thread_id\": _tid, \"new_conversation\": conversation_id is None}).encode()\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/mistral_agent\",\n"
        "        data=_data, headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=60) as _r:\n"
        "        _resp = _nimm_json.loads(_r.read())\n"
        "        return _resp.get(\"result\", \"\")\n"
    ) % tid
    ml = (
        "def nimm_mistral_list_agents(_tid='%s'):\n"
        "    _req = _nimm_ur.Request(\n"
        "        \"http://localhost:8080/api/coanimm/mistral_list_agents\",\n"
        "        data=b'{}', headers={\"Content-Type\": \"application/json\"})\n"
        "    with _nimm_ur.urlopen(_req, timeout=15) as _r:\n"
        "        return _nimm_json.loads(_r.read()).get(\"result\", \"\")\n"
    ) % tid
    parts.append(ma if "mistral_agent" not in _disabled else _stub("nimm_mistral_agent", "invoquer un agent Mistral"))
    parts.append(ml if "mistral_list_agents" not in _disabled else _stub("nimm_mistral_list_agents", "lister les agents Mistral"))
    return "".join(parts)


def _execute(code: str, args: list, workdir: str, thread_id: str = None, granted_caps=None) -> dict:
    """Écrit `code` dans un fichier temporaire de `workdir` et l'exécute.

    Retourne {'status':'ok', 'stdout':..., 'stderr':..., 'returncode':...} ou
    {'status':'error', 'message':...} en cas d'erreur de syntaxe ou délai dépassé.
    """
    syntax_err = _check_syntax(code)
    if syntax_err:
        return {'status': 'error', 'message': syntax_err,
                'stdout': '', 'stderr': '', 'returncode': 1}

    # Garde-fous de sécurité (cf. modules.coanimm_safety).
    import modules.coanimm_safety as _safety
    _risks = _safety.classify_for_execution(code)
    if _risks['blocked']:
        raisons = ' ; '.join(r['message'] for r in _risks['blocked'])
        return {'status': 'error',
                'message': f"Exécution refusée pour raison de sécurité : ce script {raisons}.",
                'blocked': _risks['blocked'],
                'stdout': '', 'stderr': '', 'returncode': 1}
    if _risks['needs_confirmation']:
        if granted_caps is None:
            # Comportement historique (run_script, exécution directe) : on bloque.
            raisons = ' ; '.join(r['message'] for r in _risks['needs_confirmation'])
            return {'status': 'error',
                    'message': (f"Ce script {raisons} : ouvre le panneau CoaNIMM pour "
                                "l'exécuter et confirmer explicitement cette action."),
                    'needs_confirmation': _risks['needs_confirmation'],
                    'stdout': '', 'stderr': '', 'returncode': 1}
        # granted_caps fourni (ex. workflow) : autorisation PAR CAPACITÉ pré-accordée.
        _caps_needed = set(_safety.capabilities_of(code)) & {'reseau', 'programme', 'email'}
        _missing = _caps_needed - set(granted_caps)
        if _missing:
            return {'status': 'error',
                    'message': ("Capacité non autorisée : " + ', '.join(sorted(_missing))
                                + ". Autorise-la dans « Capacités autorisées en exécution »."),
                    'missing_capabilities': sorted(_missing),
                    'stdout': '', 'stderr': '', 'returncode': 1}

    # Prologue = garde-fou d'écriture (confinement aux dossiers autorisés) +
    # helpers CoaNIMM (nimm_generate_image…).
    try:
        _allowed = db.list_coanimm_paths()
    except Exception:
        _allowed = []
    guard = _safety.build_guard_prologue(_allowed, allow_network=('reseau' in set(granted_caps or ())))
    prologue = _build_prologue(thread_id, workdir)
    full_code = guard + '\n' + (prologue + '\n' + code if prologue else code)
    fd, script_path = tempfile.mkstemp(suffix='.py', dir=workdir)
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as f:
            f.write(full_code)
        env = dict(os.environ)
        env['PYTHONIOENCODING'] = 'utf-8:replace'
        env['PYTHONUTF8'] = '1'
        env['PYTHONDONTWRITEBYTECODE'] = '1'  # le garde-fou bloquerait l'écriture des .pyc
        try:
            proc = subprocess.run(
                [sys.executable, script_path, *(args or [])],
                cwd=workdir,
                stdin=subprocess.DEVNULL,
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='replace',
                timeout=TIMEOUT_SECONDS,
                env=env,
            )
            return {
                'status': 'ok',
                'stdout': proc.stdout,
                'stderr': proc.stderr,
                'returncode': proc.returncode,
            }
        except subprocess.TimeoutExpired:
            return {
                'status': 'error',
                'message': f"Délai dépassé ({TIMEOUT_SECONDS}s). Le script a été interrompu.",
                'stdout': '', 'stderr': '', 'returncode': -1,
            }
    finally:
        try:
            os.unlink(script_path)
        except OSError:
            pass


def run_script(script_id: str, args: list = None, thread_id: str = None,
               confirm_scope: str = None) -> dict:
    """Exécute un script enregistré dans la Promptothèque (type='script').

    Retourne 'permission_required' si l'utilisateur doit d'abord accorder
    l'exécution (once / project / always).
    """
    action = f"exec_script:{script_id}"

    if confirm_scope in ('project', 'always'):
        db.grant_agent_permission(action, confirm_scope, thread_id)

    # La Promptothèque n'expose pas d'accès unitaire : on lit la liste des scripts
    # (type='script') et on y retrouve l'entrée. Le code est dans la clé 'text'.
    entry = db.list_prompts('script').get(script_id)

    if confirm_scope is None and not db.agent_permission_granted(action, thread_id):
        return {
            'status': 'permission_required',
            'action': action,
            'label': (entry or {}).get('label', script_id),
        }

    if not entry:
        return {'status': 'error', 'message': f"Script introuvable : {script_id}"}

    code = entry.get('text', '')
    if not code.strip():
        return {'status': 'error', 'message': "Le script est vide."}

    workdir = _workspace_dir(thread_id)
    result = _execute(code, args, workdir, thread_id)
    result['script_id'] = script_id
    return result


async def generate_code(consigne: str, thread_id: str = None,
                         provider_override: str = None) -> str:
    """Demande au LLM de générer un script Python à partir d'une consigne.

    Relance UNE fois, en demandant une version plus concise, si le premier jet est
    vide ou syntaxiquement invalide (cas typique : code coupé par max_tokens).
    Ce filet protège aussi le chemin de l'interface (/api/coanimm/generate), qui
    n'appelait pas le retry historiquement présent dans run_generated().
    """
    import core.engine as engine
    import core.hub as hub

    settings = hub.load_settings(thread_id)
    provider, model = hub.get_task_provider_model('coanimm', settings)
    if provider_override:
        provider, model = provider_override, None

    async def _ask(message: str) -> str:
        response = await engine.call_llm(
            messages=[{'role': 'user', 'content': message}],
            provider=provider,
            model=model,
            system_prompt=GENERATE_SYSTEM_PROMPT,
            max_tokens=16000,
            temperature=0.2,
            api_keys=settings['api_keys'],
        )
        return _strip_code_fences(response)

    code = await _ask(consigne)
    if code.strip() and _check_syntax(code) is None:
        return code

    # Premier jet vide ou invalide (souvent tronqué) : relance plus concise.
    print(f"[COANIMM] 1er jet invalide/tronqué, relance plus concise…")
    retry_consigne = (
        consigne
        + "\n\n[IMPORTANT : ton script précédent était invalide ou tronqué. "
        "Réécris un script Python COMPLET et plus concis : supprime les fonctions "
        "secondaires et les affichages superflus, garde l'essentiel, et assure-toi "
        "qu'il se termine proprement.]"
    )
    retry_code = await _ask(retry_consigne)
    if retry_code.strip() and _check_syntax(retry_code) is None:
        return retry_code
    # Aucune des deux tentatives n'est valide : renvoyer la plus complète des deux,
    # _check_syntax en aval produira un message d'erreur clair.
    return retry_code or code


async def repair_code(code: str, error_output: str, consigne: str = '',
                      thread_id: str = None, provider_override: str = None) -> str:
    """Corrige un script qui a échoué à l'exécution, à partir de son erreur.

    Renvoie un nouveau script Python complet, en réutilisant generate_code (donc
    avec le même nettoyage des balises et le même filet anti-troncature)."""
    objectif = (consigne or '').strip() or "(objectif initial non précisé)"
    message = (
        "Le script Python ci-dessous a échoué à l'exécution. "
        "Analyse l'erreur, corrige le script, et renvoie une version COMPLÈTE et "
        "fonctionnelle qui atteint l'objectif.\n\n"
        f"Objectif initial :\n{objectif}\n\n"
        "Script fautif :\n"
        f"{code}\n\n"
        "Sortie observée (les dernières lignes contiennent généralement l'erreur) :\n"
        f"{(error_output or '')[-2000:]}\n\n"
        "Ne réexplique pas, ne t'excuse pas : renvoie seulement le script corrigé."
    )
    return await generate_code(message, thread_id, provider_override)


def _parse_skill_fiche(raw: str) -> dict:
    """Découpe la sortie du writer (DESCRIPTION / MOTS-CLES / corps) en parties.

    Retourne {'skip': True} si le modèle a renvoyé SKIP. Sinon
    {'description': str, 'mots_cles': [..], 'corps': str}. Tolère l'absence d'une
    ligne d'en-tête (description/corps recalculés au mieux)."""
    text = (raw or '').strip()
    if not text:
        return {'skip': True}
    if text.strip().upper() == 'SKIP':
        return {'skip': True}

    description = ''
    mots_cles = []
    corps_lines = []
    header_done = False
    for line in text.splitlines():
        stripped = line.strip()
        low = stripped.lower()
        if not header_done and low.startswith('description:'):
            description = stripped.split(':', 1)[1].strip()
            continue
        if not header_done and (low.startswith('mots-cles:') or low.startswith('mots-clés:')):
            valeur = stripped.split(':', 1)[1].strip()
            mots_cles = [m.strip() for m in re.split(r'[,;]', valeur) if m.strip()]
            continue
        # Première ligne non-en-tête (et non vide) : on bascule dans le corps.
        if not header_done and not stripped:
            continue
        header_done = True
        corps_lines.append(line)

    corps = '\n'.join(corps_lines).strip()
    # Filets : si pas de description explicite, prendre la 1re phrase du corps.
    if not description and corps:
        description = corps.split('.')[0].strip()[:200]
    return {'skip': False, 'description': description,
            'mots_cles': mots_cles, 'corps': corps}


async def write_skill(consigne_origine: str, script: str, thread_id: str = None,
                      label: str = None, script_ref: str = None,
                      provider_override: str = None) -> dict:
    """Rédige et enregistre une fiche skill à partir d'un script validé par Laurent.

    Calqué sur maybe_generate_carnet_note : appel LLM en arrière-plan, lecture des
    fiches déjà écrites pour éviter un doublon, option SKIP si rien de réutilisable,
    puis écriture via save_prompt(type='skill').

    Une fiche n'existe qu'après accord explicite de Laurent : cette fonction est donc
    appelée APRÈS validation. Elle pose valide_par_laurent=True et version=1.

    Retourne {'status': 'created', 'skill': <entrée>} en cas de succès,
    {'status': 'skip'} si le modèle juge l'entrée non réutilisable, ou
    {'status': 'error', 'message': ...}.
    """
    import core.engine as engine
    import core.hub as hub

    consigne_origine = (consigne_origine or '').strip()
    script = (script or '').strip()
    if not consigne_origine and not script:
        return {'status': 'error', 'message': "Ni consigne ni script : rien à capturer."}

    settings = hub.load_settings(thread_id)
    provider, model = hub.get_task_provider_model('coanimm', settings)
    if provider_override:
        provider, model = provider_override, None

    # Fiches déjà écrites — évite de recréer un skill équivalent (cf. carnet de bord).
    existing = db.list_prompts('skill')
    existing_block = ''
    if existing:
        lignes = []
        for e in list(existing.values())[-8:]:
            desc = (e.get('meta') or {}).get('description', '') or e.get('label', '')
            if desc:
                lignes.append(f"- {desc}")
        if lignes:
            existing_block = (
                "Fiches skill déjà existantes (ne recrée pas un doublon ; réponds SKIP "
                "si ta fiche serait équivalente à l'une d'elles) :\n"
                + '\n'.join(lignes) + "\n\n"
            )

    message = (
        existing_block +
        "Consigne d'origine de Laurent :\n"
        f"{consigne_origine or '(non précisée)'}\n\n"
        "Script validé :\n"
        f"{script or '(aucun script)'}\n\n"
        "Rédige la fiche skill selon le format imposé."
    )

    try:
        raw = await engine.call_llm(
            messages=[{'role': 'user', 'content': message}],
            provider=provider,
            model=model,
            system_prompt=SKILL_WRITER_SYSTEM_PROMPT,
            max_tokens=600,
            temperature=0.3,
            api_keys=settings['api_keys'],
        )
    except Exception as e:
        detail = str(e) or type(e).__name__
        return {'status': 'error', 'message': f"Erreur génération de la fiche : {detail}"}

    fiche = _parse_skill_fiche(raw)
    if fiche.get('skip'):
        print("[SKILL] Fiche ignorée (SKIP : rien de réutilisable ou doublon).")
        return {'status': 'skip'}

    corps = fiche['corps']
    if not corps:
        return {'status': 'error', 'message': "La fiche générée est vide."}

    label = (label or fiche['description'] or 'Skill sans titre')[:120]
    import modules.coanimm_safety as _safety
    capacites = _safety.capabilities_of(script)
    meta = {
        'description': fiche['description'],
        'mots_cles': fiche['mots_cles'],
        'script_ref': script_ref or '',
        'consigne_origine': consigne_origine,
        'capacites': capacites,
        'script': script,
        'valide_par_laurent': True,
        'version': 1,
    }
    entry = db.save_prompt(None, label, corps, type='skill', meta=meta)
    print(f"[SKILL] Fiche créée : {label!r} — mots-clés : {', '.join(fiche['mots_cles']) or '(aucun)'}")
    return {'status': 'created', 'skill': entry}


def update_skill(skill_id, label=None, description=None, mots_cles=None, corps=None):
    """Met à jour un skill validé (nom, description, mots-clés, méthode) et incrémente
    sa version. Le script enregistré et les capacités sont préservés tels quels."""
    skills = db.list_prompts('skill')
    if skill_id not in skills:
        return {'status': 'error', 'message': "Skill introuvable."}
    sk = skills[skill_id]
    meta = dict(sk.get('meta') or {})
    new_label = (label if label is not None else sk.get('label', '')).strip() or sk.get('label', '')
    new_text = corps if corps is not None else sk.get('text', '')
    if description is not None:
        meta['description'] = description.strip()
    if mots_cles is not None:
        if isinstance(mots_cles, str):
            mots_cles = [m.strip() for m in mots_cles.split(',') if m.strip()]
        meta['mots_cles'] = list(mots_cles)
    try:
        meta['version'] = int(meta.get('version', 1)) + 1
    except Exception:
        meta['version'] = 2
    entry = db.save_prompt(skill_id, new_label, new_text, type='skill', meta=meta)
    print(f"[SKILL] Fiche modifiée : {new_label!r} (v{meta['version']})")
    return {'status': 'ok', 'skill': entry}


def _skill_to_text(sk: dict) -> str:
    """Met une fiche skill en texte lisible pour l'audit (label + description + corps)."""
    meta = sk.get('meta') or {}
    desc = meta.get('description', '') or sk.get('label', '')
    return (
        f"SKILL : {sk.get('label', '')}\n"
        f"Quand l'utiliser : {desc}\n"
        f"Méthode :\n{sk.get('text', '')}"
    )


def _skill_haystack(sk: dict) -> str:
    """Texte représentatif d'un skill pour l'appariement (nom + quand l'utiliser + mots-clés)."""
    meta = sk.get('meta') or {}
    return ' '.join([sk.get('label', ''), meta.get('description', ''),
                     ' '.join(meta.get('mots_cles') or [])]).strip()


def rank_skills(query: str, top_n: int = 1):
    """Classe les skills VALIDÉS par pertinence pour `query`.
    Essaie d'abord la similarité SÉMANTIQUE (embeddings « recherche par sens ») ;
    repli automatique sur le recouvrement de mots-clés si les embeddings sont
    indisponibles (modèle non installé / option désactivée). Renvoie une liste
    [(sid, sk, score)] décroissante, longueur <= top_n."""
    skills = db.list_prompts('skill')
    skills = {k: v for k, v in skills.items() if (v.get('meta') or {}).get('valide_par_laurent')}
    if not skills or not (query or '').strip():
        return []
    # 1) Sémantique (si la recherche par sens est active)
    try:
        import modules.memory as _mem
        qv = _mem._embed(query)
        if qv is not None:
            import numpy as _np
            scored = []
            for sid, sk in skills.items():
                hay = _skill_haystack(sk)
                if not hay:
                    continue
                sv = _mem._embed(hay)
                if sv is None:
                    continue
                scored.append((sid, sk, float(_np.dot(qv, sv))))
            scored = [t for t in scored if t[2] >= 0.35]
            if scored:
                scored.sort(key=lambda x: x[2], reverse=True)
                return scored[:max(1, top_n)]
    except Exception:
        pass
    # 2) Repli mots-clés
    import re as _re
    try:
        from core.hub import _MOTS_VIDES as _stop
    except Exception:
        _stop = set()
    mots = [m for m in _re.findall(r'\w+', (query or '').lower())
            if len(m) > 2 and m not in _stop]
    if not mots:
        return []
    scored = []
    for sid, sk in skills.items():
        hay = _skill_haystack(sk).lower()
        score = sum(1 for m in mots if m in hay)
        if score > 0:
            scored.append((sid, sk, score))
    scored.sort(key=lambda x: x[2], reverse=True)
    return scored[:max(1, top_n)]


def _find_relevant_skill(consigne: str):
    """Retourne la fiche skill la plus proche de la consigne (sémantique si dispo,
    sinon mots-clés), ou None. Inerte s'il n'existe aucune fiche validée."""
    res = rank_skills(consigne, top_n=1)
    return res[0][1] if res else None


def match_skills_for_consignes(consignes):
    """Pour chaque consigne, renvoie le skill VALIDÉ le plus proche (ou rien), en
    exposant l'id du skill pour composer un workflow. Renvoie une liste alignée sur
    l'entrée : {'consigne', 'skill_id', 'label', 'matched'}."""
    out = []
    for consigne in (consignes or []):
        entry = {'consigne': consigne, 'skill_id': '', 'label': '', 'matched': False}
        res = rank_skills(consigne, top_n=1)
        if res:
            sid, sk, _score = res[0]
            entry.update({'skill_id': sid, 'label': sk.get('label', ''), 'matched': True})
        out.append(entry)
    return out

async def audit_against_skill(code: str, fiche_text: str, consigne: str = '',
                              thread_id: str = None, provider_override: str = None) -> str:
    """Relit un script généré à la lumière d'une fiche skill validée et le corrige s'il
    s'écarte de la méthode décrite. Réutilise generate_code (nettoyage des balises +
    filet anti-troncature) : c'est repair_code déclenché par un écart à la fiche plutôt
    que par une erreur d'exécution. Renvoie le script (corrigé ou inchangé)."""
    objectif = (consigne or '').strip() or "(objectif initial non précisé)"
    message = (
        "Une MÉTHODE déjà validée par l'utilisateur (fiche skill) décrit comment réaliser "
        "ce type de tâche.\n\n"
        f"Fiche skill :\n{fiche_text}\n\n"
        f"Objectif :\n{objectif}\n\n"
        "Script généré :\n"
        f"{code}\n\n"
        "Le script respecte-t-il la méthode décrite dans la fiche ? S'il la respecte déjà, "
        "renvoie-le tel quel. Sinon, corrige-le pour qu'il applique cette méthode. "
        "Ne réexplique pas : renvoie seulement le script Python complet."
    )
    return await generate_code(message, thread_id, provider_override)


async def generate_plan(consigne: str, thread_id: str = None,
                         provider_override: str = None) -> dict:
    """Demande au LLM de décrire ce qu'il va faire (sans coder).
    Retourne {'plan': str, 'needs_explore': bool}."""
    import core.engine as engine
    import core.hub as hub

    settings = hub.load_settings(thread_id)
    provider, model = hub.get_task_provider_model('coanimm', settings)
    if provider_override:
        provider, model = provider_override, None
    raw = await engine.call_llm(
        messages=[{'role': 'user', 'content': consigne}],
        provider=provider,
        model=model,
        system_prompt=PLANNING_SYSTEM_PROMPT,
        max_tokens=800,
        temperature=0.3,
        api_keys=settings['api_keys'],
    )
    # Parser la ligne EXPLORER: oui/non
    needs_explore = False
    lines = (raw or '').strip().splitlines()
    plan_lines = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.lower().startswith('explorer:'):
            needs_explore = 'oui' in stripped.lower()
        else:
            plan_lines.extend(lines[i:])
            break
    return {'plan': '\n'.join(plan_lines).strip(), 'needs_explore': needs_explore}


async def explore_directory(consigne: str, thread_id: str = None,
                            confirm_scope: str = None) -> dict:
    """Génère et exécute un script d'exploration (lecture seule).
    Retourne le même format que run_generated."""
    action = GENERATED_ACTION  # permission unifiée avec l'exécution

    if confirm_scope in ('project', 'always'):
        db.grant_agent_permission(action, confirm_scope, thread_id)

    if confirm_scope is None and not db.agent_permission_granted(action, thread_id):
        return {
            'status': 'permission_required',
            'action': action,
            'label': "Explorer le disque en lecture seule",
        }

    import core.engine as engine
    import core.hub as hub

    settings = hub.load_settings(thread_id)
    provider, model = hub.get_task_provider_model('coanimm', settings)
    try:
        raw = await engine.call_llm(
            messages=[{'role': 'user', 'content': consigne}],
            provider=provider,
            model=model,
            system_prompt=EXPLORE_SYSTEM_PROMPT,
            max_tokens=16000,
            temperature=0.2,
            api_keys=settings['api_keys'],
        )
    except Exception as e:
        detail = str(e) or type(e).__name__
        return {'status': 'error',
                'message': f"Erreur lors de la génération de l'exploration : {detail}"}

    code = _strip_code_fences(raw)
    if not code.strip():
        return {'status': 'error',
                'message': "Le modèle n'a renvoyé aucun code d'exploration."}

    syntax_err = _check_syntax(code)
    if syntax_err:
        return {'status': 'error',
                'message': f"Code d'exploration invalide : {syntax_err}", 'code': code}

    workdir = _workspace_dir(thread_id)
    result = _execute(code, None, workdir, thread_id)
    result['code'] = code
    return result


def _scan_new_files(workdir: str, before: set) -> list:
    """Retourne la liste des fichiers créés dans workdir depuis le snapshot `before`."""
    try:
        after = set(os.listdir(workdir))
    except OSError:
        return []
    new = after - before
    results = []
    for fname in sorted(new):
        if fname.endswith('.py'):
            continue
        ext = os.path.splitext(fname)[1].lower()
        fpath = os.path.join(workdir, fname)
        results.append({'filename': fname, 'path': fpath, 'ext': ext})
    return results


def _route_new_files(new_files: list, thread_id: str = None) -> tuple:
    """Route les fichiers générés vers galerie ou inline.

    Retourne (info_text: str, files_list: list).
    files_list contient des dicts {filename, ext, size, url, type} pour le frontend.
    """
    if not new_files:
        return '', []
    lines = []
    files_list = []
    for f in new_files:
        fname = f['filename']
        ext   = f['ext']
        fpath = f['path']
        try:
            size = os.path.getsize(fpath)
        except OSError:
            size = 0
        tid_param = f'?thread_id={thread_id}' if thread_id else ''
        url = f'/api/coanimm/files/{fname}{tid_param}'
        if ext in _IMAGE_EXTS:
            try:
                db.save_image(fname, prompt=f'[CoaNIMM] {fname}', thread_id=thread_id or '')
                lines.append(f"Image générée et ajoutée à la galerie : {fname}")
            except Exception as e:
                lines.append(f"Image générée ({fname}) mais non sauvegardée en galerie : {e}")
            files_list.append({'filename': fname, 'ext': ext, 'size': size,
                                'url': url, 'type': 'image'})
        elif ext in _TEXT_EXTS:
            try:
                if size == 0:
                    lines.append(f"Fichier généré (vide) : {fname}")
                elif size <= _MAX_TEXT_INLINE:
                    with open(fpath, encoding='utf-8', errors='replace') as fh:
                        text_content = fh.read()
                    lines.append(f"Fichier généré : {fname}\n```\n{text_content}\n```")
                else:
                    lines.append(
                        f"Fichier généré (trop volumineux pour affichage inline, "
                        f"{size} octets) : {fname}")
            except Exception as e:
                lines.append(f"Fichier généré ({fname}) — lecture impossible : {e}")
            files_list.append({'filename': fname, 'ext': ext, 'size': size,
                                'url': url, 'type': 'text'})
        else:
            lines.append(f"Fichier généré : {fname} ({size} octets)")
            files_list.append({'filename': fname, 'ext': ext, 'size': size,
                                'url': url, 'type': 'binary'})
    return '\n'.join(lines), files_list


def execute_code(code: str, thread_id: str = None) -> dict:
    """Exécute du code Python généré par le LLM (via tool calling)."""
    workdir = _workspace_dir(thread_id)
    before  = set(os.listdir(workdir)) if os.path.isdir(workdir) else set()
    result  = _execute(code, None, workdir, thread_id)
    new_files = _scan_new_files(workdir, before)
    result['files_info'], result['files_list'] = _route_new_files(new_files, thread_id)
    result['files_count'] = len(new_files)
    return result


async def run_generated(consigne: str, thread_id: str = None,
                        confirm_scope: str = None) -> dict:
    """Génère un script Python à partir de `consigne` puis l'exécute.

    Gère automatiquement la permission exec_generated_code et relance
    une fois si le code généré est syntaxiquement invalide (tronqué).
    """
    if not consigne or not consigne.strip():
        return {'status': 'error', 'message': 'La consigne est vide.'}

    action = GENERATED_ACTION

    if confirm_scope in ('project', 'always'):
        db.grant_agent_permission(action, confirm_scope, thread_id)

    if confirm_scope is None and not db.agent_permission_granted(action, thread_id):
        return {
            'status': 'permission_required',
            'action': action,
            'label': "Génération et exécution d'un script à partir d'une consigne libre",
        }

    try:
        code = await generate_code(consigne, thread_id)
    except Exception as e:
        detail = str(e) or type(e).__name__
        print(f"[COANIMM] Erreur génération : {type(e).__name__}: {e}")
        return {'status': 'error',
                'message': f"Erreur lors de la génération du code : {detail}"}

    if not code.strip():
        return {'status': 'error', 'message': "Le modèle n'a renvoyé aucun code."}

    # Retry si syntaxe invalide (code tronqué par max_tokens)
    syntax_err = _check_syntax(code)
    if syntax_err:
        print(f"[COANIMM] Code tronqué ({syntax_err}), nouvelle tentative plus concise...")
        try:
            retry_consigne = (
                consigne
                + "\n\n[IMPORTANT : ton script précédent était trop long et a été coupé. "
                "Réécris-le de façon plus concise, en éliminant les fonctions secondaires "
                "et les affichages détaillés. L'essentiel suffit.]"
            )
            code = await generate_code(retry_consigne, thread_id)
        except Exception as e:
            detail = str(e) or type(e).__name__
            print(f"[COANIMM] Erreur regénération : {type(e).__name__}: {e}")
            return {'status': 'error',
                    'message': f"Erreur lors de la regénération : {detail}"}
        syntax_err2 = _check_syntax(code)
        if syntax_err2:
            return {
                'status': 'error',
                'message': (
                    f"Le code généré est invalide même après réécriture ({syntax_err2}). "
                    "Essaie de simplifier ta demande ou de la découper en plusieurs étapes."
                ),
                'code': code,
            }

    # Auto-audit à la lumière d'un skill validé (Étape C) — inerte si aucune fiche ne correspond.
    _fiche = _find_relevant_skill(consigne)
    if _fiche:
        try:
            _audited = await audit_against_skill(code, _skill_to_text(_fiche), consigne, thread_id)
            if _audited.strip() and _check_syntax(_audited) is None:
                code = _audited
                print("[COANIMM] Auto-audit skill appliqué avant exécution.")
        except Exception as _e:
            print(f"[COANIMM] Auto-audit skill ignoré : {_e}")

    workdir = _workspace_dir(thread_id)
    before  = set(os.listdir(workdir)) if os.path.isdir(workdir) else set()
    result  = _execute(code, None, workdir, thread_id)
    result['code'] = code
    new_files = _scan_new_files(workdir, before)
    result['files_info'], result['files_list'] = _route_new_files(new_files, thread_id)
    result['files_count'] = len(new_files)
    return result


# ══════════════════════════════════════════════════════════════════════
# WORKFLOWS — Étape 3 : séquences de skills rejouables
# ══════════════════════════════════════════════════════════════════════


def save_workflow(label: str, etapes: list, thread_id: str = None) -> dict:
    """(Stub — section workflows non disponible dans cette version.)"""
    return {'status': 'error', 'message': 'Workflows non disponibles.'}
