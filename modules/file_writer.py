# -*- coding: utf-8 -*-
"""
file_writer.py — Écriture de fichiers dans différents formats depuis le chat NIMM.

Formats supportés :
  txt, md, html, docx, pdf, epub, json, csv, daisy, daisy_audio, mp3

Tous les fichiers sont écrits dans le workspace CoaNIMM du fil courant.
Le LLM appelle l'outil write_file(filename, content, format, title, lang).
"""

import os
import re
import time
import json
import csv
import zipfile
import io
import asyncio
from pathlib import Path

# ── Formats texte brut ────────────────────────────────────────────────────────

def _write_txt(content: str, filepath: str):
    Path(filepath).write_text(content, encoding='utf-8')

def _write_md(content: str, filepath: str):
    Path(filepath).write_text(content, encoding='utf-8')

def _write_json(content: str, filepath: str):
    """Tente de parser le contenu comme JSON ; sinon wrappé dans {text: ...}."""
    try:
        data = json.loads(content)
    except Exception:
        data = {"text": content}
    Path(filepath).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')

def _write_csv(content: str, filepath: str):
    """Chaque ligne du contenu devient une ligne CSV (colonnes séparées par tabulation ou |)."""
    lines = [l for l in content.splitlines() if l.strip()]
    with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.writer(f)
        for line in lines:
            if '\t' in line:
                writer.writerow(line.split('\t'))
            elif '|' in line:
                writer.writerow([c.strip() for c in line.split('|') if c.strip()])
            else:
                writer.writerow([line])

# ── HTML ──────────────────────────────────────────────────────────────────────

def _write_html(content: str, filepath: str, title: str = 'Document', lang: str = 'fr'):
    # Si le contenu est déjà du HTML complet, on l'écrit tel quel
    if re.search(r'<html[\s>]', content, re.IGNORECASE):
        Path(filepath).write_text(content, encoding='utf-8')
        return
    # Sinon, on convertit le Markdown basique en HTML
    body = _md_to_html_body(content)
    html = f"""<!DOCTYPE html>
<html lang="{lang}">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<title>{_esc(title)}</title>
<style>
  body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 2em auto; line-height: 1.6; }}
  h1,h2,h3 {{ margin-top: 1.4em; }}
</style>
</head>
<body>
<h1>{_esc(title)}</h1>
{body}
</body>
</html>"""
    Path(filepath).write_text(html, encoding='utf-8')

def _md_to_html_body(text: str) -> str:
    """Conversion Markdown → HTML (headings, gras, listes, paragraphes)."""
    lines = text.splitlines()
    out = []
    in_ul = False
    for line in lines:
        if re.match(r'^### (.+)', line):
            if in_ul: out.append('</ul>'); in_ul = False
            out.append(f'<h3>{_esc(re.match(r"^### (.+)", line).group(1))}</h3>')
        elif re.match(r'^## (.+)', line):
            if in_ul: out.append('</ul>'); in_ul = False
            out.append(f'<h2>{_esc(re.match(r"^## (.+)", line).group(1))}</h2>')
        elif re.match(r'^# (.+)', line):
            if in_ul: out.append('</ul>'); in_ul = False
            out.append(f'<h2>{_esc(re.match(r"^# (.+)", line).group(1))}</h2>')
        elif re.match(r'^[-*] (.+)', line):
            if not in_ul: out.append('<ul>'); in_ul = True
            out.append(f'<li>{_inline_md(_esc(re.match(r"^[-*] (.+)", line).group(1)))}</li>')
        elif line.strip() == '':
            if in_ul: out.append('</ul>'); in_ul = False
            out.append('')
        else:
            if in_ul: out.append('</ul>'); in_ul = False
            out.append(f'<p>{_inline_md(_esc(line))}</p>')
    if in_ul: out.append('</ul>')
    return '\n'.join(out)

def _inline_md(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\*(.+?)\*',     r'<em>\1</em>',         text)
    text = re.sub(r'`(.+?)`',       r'<code>\1</code>',     text)
    return text

def _esc(t: str) -> str:
    return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

# ── DOCX ──────────────────────────────────────────────────────────────────────

def _write_docx(content: str, filepath: str, title: str = 'Document', lang: str = 'fr'):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.language = lang
    # Titre principal
    doc.add_heading(title, level=0)
    for line in content.splitlines():
        if re.match(r'^### (.+)', line):
            doc.add_heading(re.match(r'^### (.+)', line).group(1), level=3)
        elif re.match(r'^## (.+)', line):
            doc.add_heading(re.match(r'^## (.+)', line).group(1), level=2)
        elif re.match(r'^# (.+)', line):
            doc.add_heading(re.match(r'^# (.+)', line).group(1), level=1)
        elif re.match(r'^[-*] (.+)', line):
            p = doc.add_paragraph(re.match(r'^[-*] (.+)', line).group(1), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')
    doc.save(filepath)

# ── PDF ───────────────────────────────────────────────────────────────────────

def _write_pdf(content: str, filepath: str, title: str = 'Document', lang: str = 'fr'):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_LEFT

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        title=title,
        author='NIMM',
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=16, spaceAfter=12)
    h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=13, spaceAfter=8)
    h3 = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=11, spaceAfter=6)
    body = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11, leading=15, spaceAfter=6)

    story = [Paragraph(title, h1), Spacer(1, 0.4*cm)]
    for line in content.splitlines():
        if re.match(r'^### (.+)', line):
            story.append(Paragraph(re.match(r'^### (.+)', line).group(1), h3))
        elif re.match(r'^## (.+)', line):
            story.append(Paragraph(re.match(r'^## (.+)', line).group(1), h2))
        elif re.match(r'^# (.+)', line):
            story.append(Paragraph(re.match(r'^# (.+)', line).group(1), h1))
        elif line.strip():
            story.append(Paragraph(line.replace('**','<b>',1).replace('**','</b>',1), body))
        else:
            story.append(Spacer(1, 0.3*cm))
    doc.build(story)

# ── EPUB ──────────────────────────────────────────────────────────────────────

def _write_epub(content: str, filepath: str, title: str = 'Document', lang: str = 'fr'):
    """Génère un EPUB 3 minimaliste via structure zip manuelle."""
    uid = f"nimm-{int(time.time())}"
    body_html = _md_to_html_body(content)

    content_doc = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="{lang}" lang="{lang}">
<head><meta charset="UTF-8"/><title>{_esc(title)}</title></head>
<body>
<h1>{_esc(title)}</h1>
{body_html}
</body>
</html>"""

    opf = f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="uid" xml:lang="{lang}">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="uid">{uid}</dc:identifier>
    <dc:title>{_esc(title)}</dc:title>
    <dc:language>{lang}</dc:language>
  </metadata>
  <manifest>
    <item id="content" href="content.xhtml" media-type="application/xhtml+xml"/>
    <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
  </manifest>
  <spine toc="ncx">
    <itemref idref="content"/>
  </spine>
</package>"""

    ncx = f"""<?xml version="1.0" encoding="UTF-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head><meta name="dtb:uid" content="{uid}"/></head>
  <docTitle><text>{_esc(title)}</text></docTitle>
  <navMap>
    <navPoint id="np1" playOrder="1">
      <navLabel><text>{_esc(title)}</text></navLabel>
      <content src="content.xhtml"/>
    </navPoint>
  </navMap>
</ncx>"""

    with zipfile.ZipFile(filepath, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('mimetype', 'application/epub+zip', compress_type=zipfile.ZIP_STORED)
        zf.writestr('META-INF/container.xml',
            '<?xml version="1.0"?>'
            '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
            '<rootfiles><rootfile full-path="OEBPS/content.opf"'
            ' media-type="application/oebps-package+xml"/></rootfiles></container>')
        zf.writestr('OEBPS/content.opf', opf)
        zf.writestr('OEBPS/content.xhtml', content_doc)
        zf.writestr('OEBPS/toc.ncx', ncx)

# ── DAISY 2.02 (texte seul, sans audio) ──────────────────────────────────────

def _write_daisy_text(content: str, dirpath: str, title: str = 'Document', lang: str = 'fr'):
    """
    Génère un livre DAISY 2.02 texte seul dans un répertoire.
    Structure : ncc.html + content_001.html + content_001.smil (silences).
    Compatible avec les lecteurs DAISY (AMIS, Victor Reader, Dolphin EasyReader…).
    """
    os.makedirs(dirpath, exist_ok=True)
    uid = f"nimm-{int(time.time())}"

    # Découper le contenu en sections (headings H1/H2)
    sections = _split_sections(content, title)

    # ── content HTML par section ──────────────────────────────────────────
    smil_refs = []
    for i, (sec_title, paras) in enumerate(sections, 1):
        sec_id = f"s{i:03d}"
        html_name = f"content_{i:03d}.html"
        smil_name = f"content_{i:03d}.smil"

        # Générer les paragraphes avec ID
        para_items = []
        for j, para in enumerate(paras, 1):
            pid = f"{sec_id}_p{j}"
            para_items.append((pid, para))

        # HTML de la section
        paras_html = '\n'.join(
            f'<p id="{pid}">{_esc(p)}</p>' for pid, p in para_items
        )
        sec_html = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN"
  "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd">
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="{lang}" lang="{lang}">
<head><meta http-equiv="Content-Type" content="text/html; charset=UTF-8"/>
<title>{_esc(sec_title)}</title></head>
<body>
<h1 id="{sec_id}">{_esc(sec_title)}</h1>
{paras_html}
</body></html>"""
        (Path(dirpath) / html_name).write_text(sec_html, encoding='utf-8')

        # SMIL de la section (texte seul — pas de pointeurs audio)
        smil_seqs = '\n'.join(
            f'    <par id="par_{pid}"><text src="{html_name}#{pid}"/></par>'
            for pid, _ in para_items
        )
        dur_total = sum(max(1, len(p)//15) for _, p in para_items)  # estimation
        smil = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE smil PUBLIC "-//W3C//DTD SMIL 1.0//EN"
  "http://www.w3.org/TR/REC-smil/SMIL10.dtd">
<smil>
<head>
  <meta name="ncc:timeInThisSmil" content="{dur_total}s"/>
  <meta name="title" content="{_esc(sec_title)}"/>
  <layout><region id="textView"/></layout>
</head>
<body>
  <seq id="sq{i:03d}" dur="{dur_total}s">
    <par id="par_{sec_id}"><text src="{html_name}#{sec_id}"/></par>
{smil_seqs}
  </seq>
</body></smil>"""
        (Path(dirpath) / smil_name).write_text(smil, encoding='utf-8')
        smil_refs.append((sec_id, sec_title, smil_name, html_name, dur_total))

    # ── NCC.html ──────────────────────────────────────────────────────────
    total_dur = sum(d for _, _, _, _, d in smil_refs)
    nav_items = '\n'.join(
        f'  <h1 class="title" id="{sid}"><a href="{smil}#{sid}">{_esc(stitle)}</a></h1>'
        for sid, stitle, smil, _, _ in smil_refs
    )
    ncc = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN"
  "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd">
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="{lang}" lang="{lang}">
<head>
  <meta http-equiv="Content-Type" content="text/html; charset=UTF-8"/>
  <meta name="dc:title" content="{_esc(title)}"/>
  <meta name="dc:language" content="{lang}"/>
  <meta name="dc:format" content="Daisy 2.02"/>
  <meta name="dc:type" content="Text"/>
  <meta name="ncc:totalTime" content="{total_dur}s"/>
  <meta name="ncc:tocItems" content="{len(smil_refs)}"/>
  <meta name="ncc:depth" content="1"/>
  <meta name="ncc:pageNormal" content="0"/>
  <meta name="ncc:generator" content="NIMM file_writer"/>
  <title>{_esc(title)}</title>
</head>
<body>
{nav_items}
</body></html>"""
    (Path(dirpath) / 'ncc.html').write_text(ncc, encoding='utf-8')

def _split_sections(content: str, default_title: str) -> list:
    """Découpe le contenu en [(titre, [paragraphes])]."""
    sections = []
    current_title = default_title
    current_paras = []
    for line in content.splitlines():
        m = re.match(r'^#{1,2} (.+)', line)
        if m:
            if current_paras:
                sections.append((current_title, [p for p in current_paras if p.strip()]))
            current_title = m.group(1).strip()
            current_paras = []
        elif line.strip():
            # Nettoyer le Markdown inline
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            clean = re.sub(r'^[-*] ', '', clean)
            current_paras.append(clean.strip())
    if current_paras:
        sections.append((current_title, [p for p in current_paras if p.strip()]))
    if not sections:
        sections = [(default_title, [content.strip()])]
    return sections

# ── DAISY 2.02 avec audio (edge-tts + lameenc) ───────────────────────────────

async def _write_daisy_audio_async(content: str, dirpath: str,
                                    title: str = 'Document', lang: str = 'fr',
                                    voice: str = ''):
    """
    Génère un livre DAISY 2.02 avec audio MP3 via edge-tts.
    Un fichier MP3 par section, synchronisation SMIL texte+audio.
    """
    import edge_tts
    import lameenc

    os.makedirs(dirpath, exist_ok=True)
    uid = f"nimm-{int(time.time())}"

    # Choisir la voix selon la langue
    if not voice:
        voice_map = {
            'fr': 'fr-FR-DeniseNeural',
            'en': 'en-GB-SoniaNeural',
            'es': 'es-ES-ElviraNeural',
            'de': 'de-DE-KatjaNeural',
        }
        voice = voice_map.get(lang[:2], 'fr-FR-DeniseNeural')

    sections = _split_sections(content, title)
    smil_refs = []

    for i, (sec_title, paras) in enumerate(sections, 1):
        sec_id  = f"s{i:03d}"
        html_name = f"content_{i:03d}.html"
        smil_name = f"content_{i:03d}.smil"
        mp3_name  = f"audio_{i:03d}.mp3"

        # Générer audio via edge-tts
        full_text = f"{sec_title}. " + " ".join(paras)
        communicate = edge_tts.Communicate(full_text, voice)
        mp3_path = os.path.join(dirpath, mp3_name)
        await communicate.save(mp3_path)

        # Durée approx (edge-tts → MP3, ~150 mots/min, 1 mot ≈ 5 chars)
        dur_s = max(2, len(full_text) // 12)

        # Paragraphes avec IDs
        para_items = [(f"{sec_id}_p{j}", p) for j, p in enumerate(paras, 1)]

        # HTML de la section
        paras_html = '\n'.join(
            f'<p id="{pid}">{_esc(p)}</p>' for pid, p in para_items
        )
        sec_html = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN"
  "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd">
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="{lang}" lang="{lang}">
<head><meta http-equiv="Content-Type" content="text/html; charset=UTF-8"/>
<title>{_esc(sec_title)}</title></head>
<body>
<h1 id="{sec_id}">{_esc(sec_title)}</h1>
{paras_html}
</body></html>"""
        (Path(dirpath) / html_name).write_text(sec_html, encoding='utf-8')

        # SMIL avec pointeur audio
        clip_end = dur_s
        smil_seq_items = '\n'.join(
            f'    <par id="par_{pid}"><text src="{html_name}#{pid}"/>'
            f'<audio src="{mp3_name}" clip-begin="npt=0s" clip-end="npt={clip_end}s"/></par>'
            for pid, _ in para_items
        )
        smil = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE smil PUBLIC "-//W3C//DTD SMIL 1.0//EN"
  "http://www.w3.org/TR/REC-smil/SMIL10.dtd">
<smil>
<head>
  <meta name="ncc:timeInThisSmil" content="{dur_s}s"/>
  <meta name="title" content="{_esc(sec_title)}"/>
  <layout><region id="textView"/></layout>
</head>
<body>
  <seq id="sq{i:03d}" dur="{dur_s}s">
    <par id="par_{sec_id}">
      <text src="{html_name}#{sec_id}"/>
      <audio src="{mp3_name}" clip-begin="npt=0s" clip-end="npt={dur_s}s"/>
    </par>
{smil_seq_items}
  </seq>
</body></smil>"""
        (Path(dirpath) / smil_name).write_text(smil, encoding='utf-8')
        smil_refs.append((sec_id, sec_title, smil_name, html_name, dur_s))

    # NCC.html
    total_dur = sum(d for _, _, _, _, d in smil_refs)
    nav_items = '\n'.join(
        f'  <h1 class="title" id="{sid}"><a href="{smil}#{sid}">{_esc(st)}</a></h1>'
        for sid, st, smil, _, _ in smil_refs
    )
    ncc = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN"
  "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd">
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="{lang}" lang="{lang}">
<head>
  <meta http-equiv="Content-Type" content="text/html; charset=UTF-8"/>
  <meta name="dc:title" content="{_esc(title)}"/>
  <meta name="dc:language" content="{lang}"/>
  <meta name="dc:format" content="Daisy 2.02"/>
  <meta name="dc:type" content="Text_with_audio"/>
  <meta name="ncc:totalTime" content="{total_dur}s"/>
  <meta name="ncc:tocItems" content="{len(smil_refs)}"/>
  <meta name="ncc:depth" content="1"/>
  <meta name="ncc:pageNormal" content="0"/>
  <meta name="ncc:generator" content="NIMM file_writer"/>
  <title>{_esc(title)}</title>
</head>
<body>
{nav_items}
</body></html>"""
    (Path(dirpath) / 'ncc.html').write_text(ncc, encoding='utf-8')

    # Zipper le tout en un .daisy
    zip_path = dirpath.rstrip('/\\') + '.daisy'
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for f in os.listdir(dirpath):
            zf.write(os.path.join(dirpath, f), f)
    return zip_path

# ── MP3 ───────────────────────────────────────────────────────────────────────

async def _write_mp3_async(content: str, filepath: str,
                            title: str = 'Document', lang: str = 'fr',
                            voice: str = ''):
    """Génère un fichier MP3 via edge-tts."""
    import edge_tts
    if not voice:
        voice_map = {
            'fr': 'fr-FR-DeniseNeural',
            'en': 'en-GB-SoniaNeural',
            'es': 'es-ES-ElviraNeural',
            'de': 'de-DE-KatjaNeural',
        }
        voice = voice_map.get(lang[:2], 'fr-FR-DeniseNeural')
    communicate = edge_tts.Communicate(content, voice)
    await communicate.save(filepath)

# ── Point d'entrée principal ──────────────────────────────────────────────────

SUPPORTED_FORMATS = {
    'txt':         '.txt',
    'md':          '.md',
    'markdown':    '.md',
    'html':        '.html',
    'docx':        '.docx',
    'word':        '.docx',
    'pdf':         '.pdf',
    'epub':        '.epub',
    'json':        '.json',
    'csv':         '.csv',
    'daisy':       '.daisy',   # texte seul (répertoire zippé)
    'daisy_audio': '.daisy',   # texte + audio MP3
    'mp3':         '.mp3',
    'audio':       '.mp3',
}

def write_file(content: str, fmt: str, filepath: str,
               title: str = 'Document', lang: str = 'fr',
               voice: str = '') -> str:
    """
    Écrit le contenu dans le format demandé.

    Args:
        content : texte source (Markdown ou plain text)
        fmt     : format cible (voir SUPPORTED_FORMATS)
        filepath: chemin complet du fichier de sortie (sans extension si daisy)
        title   : titre du document
        lang    : code langue ISO 639-1 ('fr', 'en', …)
        voice   : voix edge-tts (optionnel, sinon choisie selon lang)

    Returns:
        Chemin réel du fichier créé.
    """
    fmt = fmt.lower().strip()
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"Format non supporté : {fmt}. "
                         f"Formats disponibles : {', '.join(SUPPORTED_FORMATS)}")

    # Assurer l'extension correcte
    ext = SUPPORTED_FORMATS[fmt]
    if not filepath.endswith(ext):
        filepath = re.sub(r'\.[^.]+$', '', filepath) + ext

    os.makedirs(os.path.dirname(filepath) or '.', exist_ok=True)

    if fmt in ('txt',):
        _write_txt(content, filepath)
    elif fmt in ('md', 'markdown'):
        _write_md(content, filepath)
    elif fmt == 'json':
        _write_json(content, filepath)
    elif fmt == 'csv':
        _write_csv(content, filepath)
    elif fmt == 'html':
        _write_html(content, filepath, title, lang)
    elif fmt in ('docx', 'word'):
        _write_docx(content, filepath, title, lang)
    elif fmt == 'pdf':
        _write_pdf(content, filepath, title, lang)
    elif fmt == 'epub':
        _write_epub(content, filepath, title, lang)
    elif fmt == 'daisy':
        dirpath = filepath.replace('.daisy', '_daisy_pkg')
        _write_daisy_text(content, dirpath, title, lang)
        # Zipper le répertoire DAISY
        with zipfile.ZipFile(filepath, 'w', zipfile.ZIP_DEFLATED) as zf:
            for f in os.listdir(dirpath):
                zf.write(os.path.join(dirpath, f), f)
    elif fmt == 'daisy_audio':
        dirpath = filepath.replace('.daisy', '_daisy_audio_pkg')
        result = asyncio.run(_write_daisy_audio_async(content, dirpath, title, lang, voice))
        filepath = result
    elif fmt in ('mp3', 'audio'):
        asyncio.run(_write_mp3_async(content, filepath, title, lang, voice))

    return filepath
